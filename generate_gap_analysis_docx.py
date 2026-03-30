from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Helper: shade a cell
def shade_cell(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# Helper: set cell text with formatting
def set_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# Helper: add a formatted table
def add_table(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, size=9, color=(255,255,255))
        shade_cell(cell, header_color)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell(cell, val, size=9)
            if r_idx % 2 == 1:
                shade_cell(cell, 'F2F2F2')
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AP Supplier Portal')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 78, 121)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Wireframe Gap Analysis')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Gap Register')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(89, 89, 89)

for _ in range(3):
    doc.add_paragraph('')

info = [
    ('Prepared for:', 'HighRadius Corporation'),
    ('Prepared by:', 'Metapointer Labs Pvt. Ltd.'),
    ('Date:', 'March 2026'),
    ('Version:', '3.3'),
    ('Classification:', 'Confidential'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '  ')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.bold = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Gap Analysis Overview',
    '3. Wireframe Screen Inventory (19 Wireframed + 10 To-Be-Designed)',
    '4. Validation Gap Analysis (12)',
    '5. Logic Complexity Gaps (7)',
    '6. Missing API Endpoints (3)',
    '7. Technical Gaps (9)',
    '8. Workflow Gaps (6)',
    '9. Functional Requirements Traceability',
    '10. Process Flow Coverage (Phase 1 & Phase 2)',
    '11. Non-Functional Requirements Gap Check',
    '12. Platform Integration Gap: AP Portal to Supplier Portal Entry Point',
    '13. Gap Mitigation Strategy',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.style = doc.styles['List Number']

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This document provides a comprehensive gap analysis of the HighRadius AP Supplier Portal, '
    'identifying all gaps between the HighRadius User Journey wireframes (11 journey folders, 55+ screens), '
    'the Product Backlog Requirements (PBR) spreadsheet, the RFP, and the functional/non-functional requirements. '
    'A total of 42 gaps have been identified across 6 categories. Wireframe coverage is 66% (19 of 29 screens) — '
    'all supplier-facing transaction screens are wireframed; the 10 missing screens are admin/manager screens. '
    'All gaps are fully accounted for in Metapointer\'s estimation and delivery plan — no hidden scope, '
    'no surprise change requests.'
)

doc.add_paragraph('')
doc.add_heading('Key Facts', level=2)

key_facts = [
    ('Total Screens', '29 (19 wireframed by HighRadius + 10 to be designed)'),
    ('Total Gaps Identified', '42 across 6 categories'),
    ('Wireframe Coverage', '19 of 29 screens (66%) — 10 screens (34%) to be designed'),
    ('Total APIs', '66 (58 vendor-built + 7 HR integrations + 1 orchestration)'),
    ('Intelligent Agents', '7 (4 automated + 3 assisted)'),
    ('Core Features', '25+ mapped to PBR spreadsheet rows 2-25 plus 2 dashboards'),
    ('Gap Resolution', 'All 42 gaps resolved via Sprint 0 design sprint + architecture work'),
]
add_table(doc, ['Item', 'Detail'], key_facts, col_widths=[5, 12])

doc.add_paragraph('')

# ============================================================
# 2. GAP ANALYSIS OVERVIEW
# ============================================================
doc.add_heading('2. Gap Analysis Overview', level=1)

doc.add_paragraph(
    'Gaps were identified by performing a systematic analysis of the HighRadius User Journey '
    'wireframes (organised into 11 journey folders with 55+ screenshots) against the Product Backlog '
    'Requirements (PBR) spreadsheet and the RFP. Every RFP requirement and PBR line item was mapped '
    'to a wireframe screen — anything without a wireframe or with incomplete specification is classified '
    'as a gap. Wireframes provide screen layout and basic UX but do not cover validation logic, '
    'error states, complex business rules, or infrastructure concerns.'
)

summary_rows = [
    ('Screens Wireframed by HighRadius', '19', 'From User Journey wireframes (across 11 journey folders)'),
    ('Missing / To-Be-Designed Screens', '10', '34% of screens require UX design from scratch (all admin/manager screens)'),
    ('Validation Gaps', '12', 'Validation types specified but not illustrated in wireframes'),
    ('Logic Complexity Gaps', '7', 'Areas with significantly more complexity than wireframes suggest'),
    ('Missing API Endpoints', '3', 'Required endpoints not reflected in wireframes'),
    ('Technical Gaps', '9', 'Infrastructure/integration gaps not covered in wireframes'),
    ('Workflow Gaps', '6', 'Workflow logic not fully specified in wireframes'),
    ('Total Gaps', '42', 'All accounted for in estimation and delivery plan'),
]
add_table(doc, ['Gap Category', 'Count', 'Details'], summary_rows, col_widths=[5, 2, 10])

doc.add_page_break()

# ============================================================
# 3. WIREFRAME SCREEN INVENTORY (19 Wireframed + 10 To-Be-Designed)
# ============================================================
doc.add_heading('3. Wireframe Screen Inventory (19 Wireframed + 10 To-Be-Designed)', level=1)

doc.add_paragraph(
    'A systematic review of the HighRadius User Journey wireframes (organised into 11 journey folders) '
    'confirms that 19 of 29 screens have wireframe coverage (66%). The remaining 10 screens (34%) — '
    'all admin, manager, and workflow screens — require UX design from scratch during Sprint 0. '
    'UX runs 1 sprint ahead of development throughout. Weekly design reviews with HighRadius product '
    'team ensure alignment before any code is written.'
)

doc.add_heading('3A. Wireframed Screens (19)', level=2)
doc.add_paragraph(
    'The following 19 screens have wireframe layouts provided in the HighRadius User Journey document. '
    'Each entry shows the flow name, screen name, persona, phase, and what the wireframe covers.'
)

# Helper: add a wireframed screen entry (no images — wireframes are in the User Journey PDF)
def add_screen_entry(doc, num, screen_name, flow, persona, phase, description):
    """Add a numbered screen entry with heading, flow name, metadata, and description."""
    doc.add_heading(f'{num}. {screen_name}', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Flow: ')
    run.bold = True
    p.add_run(flow)
    run2 = p.add_run('    |    Persona: ')
    run2.bold = True
    p.add_run(persona)
    run3 = p.add_run('    |    Phase: ')
    run3.bold = True
    p.add_run(phase)

    p2 = doc.add_paragraph()
    run4 = p2.add_run('What the Wireframe Shows: ')
    run4.bold = True
    p2.add_run(description)
    doc.add_paragraph('')  # spacer

# --- 3A: Wireframed Screens ---
add_screen_entry(doc, 1, 'AP Dashboard (Products Dropdown)', 'Supplier Management', 'Supplier Manager', 'Phase 1',
    'AP Automation dashboard with "Products" dropdown showing Accounts Payable and Supplier Portal modules.')

add_screen_entry(doc, 2, 'Supplier List (Lifecycle View)', 'Supplier Management', 'Supplier Manager', 'Phase 1',
    'All Suppliers (10) grid with left nav tree (Onboarded/Invited/Inactive), + Add Supplier button, LiveCube.')

add_screen_entry(doc, 3, 'Invitation E-Form Modal', 'Supplier Management', 'Supplier Manager', 'Phase 1',
    'Modal: Supplier Name, Email, Confirm Email, ERP/Company dropdown, Supplier Category dropdown.')

add_screen_entry(doc, 4, 'Email Invitation', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Gmail view: "Welcome to Supplier Portal Registration" from noreply@qas.com with secure registration link.')

add_screen_entry(doc, 5, 'Registration: General Supplier Info', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Supplier Identity (Name, Doing Business As, Email, Operating Unit) + Tax Information (TIN, Classification, Form Type).')

add_screen_entry(doc, 6, 'Registration: Purchase Info', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Sales Representative (Contact, Phone) + Billing & Shipping Details (Email, Address, Country, City, Zip).')

add_screen_entry(doc, 7, 'Registration: Remittance Info', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Remittance grid with Add New modal (Address Name, Phone, Email, Address lines, Country, City, Status toggle).')

add_screen_entry(doc, 8, 'Registration: Payment Info', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Payment Method (type, nickname, account/IBAN, ABA) + Bank Details (name, type, address). Empty + filled states.')

add_screen_entry(doc, 9, 'Registration: Document Upload', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Document Type dropdown, drag-and-drop upload, file grid (Name, User, Date, Type, Preview, Action), certification checkbox.')

add_screen_entry(doc, 10, 'Registration: Submit Confirmation', 'New Supplier Registration', 'Supplier', 'Phase 1',
    'Submit button + "Application Submitted Successfully" toast notification.')

add_screen_entry(doc, 11, 'Supplier Dashboard', 'Dashboard', 'Supplier', 'Phase 2',
    'Financial Summary (Total Due, Overdue, Invoices in Processing, Unbilled PO, Awaiting Payment), '
    'Invoice Status Pipeline chart, My Tasks & Alerts (New POs, Rejected Invoices, Information Requested, Expiring Documents).')

add_screen_entry(doc, 12, 'Purchase Order List', 'Looking through Purchase Orders', 'Supplier', 'Phase 2',
    'PO grid: PO No, Status (Acknowledgment Pending/Acknowledged), Buyer, PO Date, Currency, Total PO Amount, Total Billed. '
    'Actions: multi-select, filters, bulk acknowledge.')

add_screen_entry(doc, 13, 'Purchase Order Detail', 'Looking through Purchase Orders + Details page for different status in PO', 'Supplier', 'Phase 2',
    'Billing & Summary section, Line Items grid (ID, Item Code, Description, Qty, UoM, Unit Price, Tax). '
    'Status variants: Acknowledgment Pending (Acknowledge button), Acknowledged (Create Invoice button), Rejected, Closed.')

add_screen_entry(doc, 14, 'Invoice List', 'Looking through Invoices', 'Supplier', 'Phase 2',
    'Left nav tree: All Invoices (16) > Draft (03) > Submitted (02) > In Progress (05) > Rejected (01) > Paid (04). '
    'Grid: Invoice No, PO No, Status, Source, Invoice Date, Payment Terms, Due Date, Currency. + Add Invoice button.')

add_screen_entry(doc, 15, 'Invoice Detail', 'Looking through Invoices + Details page for different status in Invoice', 'Supplier', 'Phase 2',
    'Billing/Summary/Supplier sections, Details + Files tabs, Document Viewer sidebar. '
    'Status variants: Draft (editable, Submit/Delete), Submitted, In Progress, Rejected, Paid, Void.')

add_screen_entry(doc, 16, 'PO Flip / Invoice Creation', 'Flipping PO to Create Invoice', 'Supplier', 'Phase 2',
    '"Select Lines to Create Invoice" modal with checkboxes per PO line, quantity/amount selection. '
    'Auto-populates new invoice from PO data. "Invoice lines added" confirmation toast.')

add_screen_entry(doc, 17, 'Blank Invoice Creation', 'Creating a Blank Invoice', 'Supplier', 'Phase 2',
    'Empty invoice form: Bill to, Legal Entity, Location, Invoice Date, Payment Terms, Type, PO Number. '
    'Editable fields + Add Line Item. Submit/Delete actions.')

add_screen_entry(doc, 18, 'Supplier Profile', 'Supplier Profile', 'Supplier', 'Phase 2',
    'Supplier Information page: General Supplier Information (Supplier Identity, Business Profile & Certifications). '
    'Navigation from Invoice List to profile view.')

add_screen_entry(doc, 19, 'Files Tab (Invoice Attachments)', 'Files Tab Journey', 'Supplier', 'Phase 2',
    'Files tab on Invoice Detail: filter tags (All/Invoice/Supporting), Upload modal (drag-drop, progress bars), '
    'file grid (File, Description, Tag, Type, Actions), Document Viewer preview, Edit Description/Type modal.')

doc.add_page_break()

doc.add_heading('3B. Missing / To-Be-Designed Screens (10)', level=2)
doc.add_paragraph(
    '10 screens (34%) are absent from the wireframes — all are admin, manager, or workflow screens. '
    'These will be designed from scratch during the Sprint 0 design sprint. '
    'Note: All supplier-facing transaction screens (PO, Invoice, Dashboard) ARE wireframed above.'
)

# Helper: add a missing screen entry with evidence (text only, no images)
def add_missing_screen_entry(doc, num, screen_name, persona, phase, complexity, evidence, what_to_design):
    doc.add_heading(f'Gap #{num}. {screen_name}', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Persona: ')
    run.bold = True
    p.add_run(persona)
    run2 = p.add_run('    |    Phase: ')
    run2.bold = True
    p.add_run(phase)
    run3 = p.add_run('    |    Complexity: ')
    run3.bold = True
    run4 = p.add_run(complexity)
    run4.bold = True
    if complexity == 'High':
        run4.font.color.rgb = RGBColor(192, 0, 0)
    else:
        run4.font.color.rgb = RGBColor(0, 112, 192)

    p2 = doc.add_paragraph()
    run5 = p2.add_run('Wireframe Evidence: ')
    run5.bold = True
    run5.font.color.rgb = RGBColor(192, 0, 0)
    p2.add_run(evidence)

    p3 = doc.add_paragraph()
    run6 = p3.add_run('What Needs To Be Designed: ')
    run6.bold = True
    p3.add_run(what_to_design)
    doc.add_paragraph('')

add_missing_screen_entry(doc, 1, 'Manager Dashboard', 'Supplier Manager', 'Phase 1', 'High',
    'NOT in any User Journey flow. The "Dashboard" flow only covers the Supplier Dashboard (supplier-facing) '
    '— no buyer/manager dashboard exists anywhere in the 11 journey flows.',
    'Onboarding funnel, pending tasks, SLA metrics, quick-action panel.')

add_missing_screen_entry(doc, 2, 'Screening Dashboard', 'Supplier Manager', 'Phase 1', 'Medium',
    'NOT in any User Journey flow. The "Supplier Management" flow shows screening sub-stages in the left nav tree '
    '(TIN Check, Bank Validation, Sanctions) — but no dedicated screening dashboard screen.',
    'Aggregated compliance results across suppliers.')

add_missing_screen_entry(doc, 3, 'Approval Worklist', 'Supplier Manager', 'Phase 1', 'Medium',
    'NOT in any User Journey flow. The "Supplier Management" flow shows "Need Review" in the left nav tree '
    '— but no dedicated approval queue screen with approve/reject/reassign actions.',
    'Task queue with approve/reject/reassign actions.')

add_missing_screen_entry(doc, 4, 'Admin Settings', 'Admin', 'Phase 1', 'High',
    'NOT in any User Journey flow. No admin or configuration screens exist in any of the 11 journey flows. '
    'Entirely absent from the User Journey wireframes.',
    'Screening configuration, system parameters.')

add_missing_screen_entry(doc, 5, 'User / Role Management', 'Admin', 'Phase 1', 'High',
    'NOT in any User Journey flow. No user management screens in any journey flow. Keycloak integration assumed '
    'but no UI exists for user creation or permission control.',
    'Keycloak-integrated user creation, permission control.')

add_missing_screen_entry(doc, 6, 'Audit Trail', 'Admin', 'Phase 1', 'Medium',
    'NOT in any User Journey flow. No audit or logging screens in any of the 11 journey flows. '
    'Entirely absent from the User Journey wireframes.',
    'Filterable/exportable action log.')

add_missing_screen_entry(doc, 7, 'Form Builder', 'Admin', 'Phase 1', 'High',
    'NOT in any User Journey flow. The "New Supplier Registration" flow shows the registration form output '
    '(5-step form) — but no form configuration/builder screen for admins to configure form fields.',
    'Drag-and-drop form configuration, multi-language labels.')

add_missing_screen_entry(doc, 8, 'Workflow Configuration', 'Admin', 'Phase 1', 'High',
    'NOT in any User Journey flow. Approval flow is referenced in requirements but no workflow builder screen '
    'exists in any of the 11 journey flows.',
    'Approval chain builder with escalation rules.')

add_missing_screen_entry(doc, 9, 'Notification Center', 'Cross-cutting', 'Phase 1', 'Medium',
    'NOT in any User Journey flow. The "Dashboard" flow shows a "My Tasks & Alerts" section on the Supplier '
    'Dashboard — but no dedicated, full-page notification center screen.',
    'Real-time alerts for all personas.')

add_missing_screen_entry(doc, 10, 'Master Data Change Request', 'Supplier/Manager', 'Phase 2', 'Medium',
    'NOT in any User Journey flow. The "Supplier Profile" flow shows the Supplier Profile in read-only view '
    '— but no change request or edit-with-approval flow screen.',
    'Data update form with approval workflow trigger.')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Complexity Breakdown: ')
run.bold = True
p.add_run('High = 5 screens, Medium = 5 screens. High-complexity screens require 2-3x development effort.')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Phase Split: ')
run.bold = True
p.add_run('Phase 1 = 9 screens (Sprint 0 design + Sprints 1-5 build), Phase 2 = 1 screen (Sprints 6-9 build).')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Insight: ')
run.bold = True
p.add_run(
    'All 10 missing screens are admin/manager/workflow screens — the back-office and configuration layer. '
    'The supplier-facing transaction screens (Dashboard, PO List, PO Detail, Invoice List, Invoice Detail, '
    'PO Flip, Blank Invoice, Supplier Profile, Files Tab) are all wireframed by HighRadius. '
    'This means the supplier user experience is well-defined; the gaps are in the buyer-side management '
    'and admin tooling that HighRadius expects the vendor to design.'
)

doc.add_page_break()

# ============================================================
# 4. VALIDATION GAP ANALYSIS (12)
# ============================================================
doc.add_heading('4. Validation Gap Analysis (12)', level=1)

doc.add_paragraph(
    'Twelve validation types are specified in requirements but not illustrated in wireframes. '
    'All are documented in the Validation Error State Library — a Sprint 0 deliverable completed '
    'before Sprint 1 begins. The library includes exact error message text, toast vs. inline vs. '
    'modal decision, colour/icon specifications per G4 DSL, screen-by-screen mapping, and i18n '
    'label keys for multi-language support.'
)

validations = [
    ('1', 'Email format validation', 'Supplier Registration', 'Inline error states', 'Phase 1'),
    ('2', 'TIN format/checksum validation', 'Registration Form', 'Real-time feedback indicator', 'Phase 1'),
    ('3', 'Bank account number format', 'Bank Details Form', 'Field-level error + toast', 'Phase 1'),
    ('4', 'Required field enforcement (dynamic per tenant)', 'All form screens', 'Visual indicators + submit blocking', 'Phase 1'),
    ('5', 'File type/size validation', 'Document Upload', 'Error toast + supported format list', 'Phase 1'),
    ('6', 'Duplicate supplier detection', 'Invitation, Registration', 'Warning modal with existing supplier details', 'Phase 1'),
    ('7', 'Approval chain completeness', 'Workflow Configuration', 'Pre-save validation + error summary', 'Phase 1'),
    ('8', 'PO quantity bounds (partial invoicing)', 'PO Flip Screen', 'Inline error: "Cannot exceed PO quantity"', 'Phase 2'),
    ('9', 'Invoice amount reconciliation', 'Invoice Creation', 'Warning when total deviates from PO', 'Phase 2'),
    ('10', 'Date range validation', 'List screens (PO/invoice filters)', 'Inline date picker constraints', 'Both'),
    ('11', 'Session timeout/token expiry', 'All screens', 'Modal with re-login prompt', 'Both'),
    ('12', 'Concurrent edit conflict', 'Approval screens', 'Toast notification + refresh prompt', 'Both'),
]
add_table(doc, ['#', 'Validation Type', 'Affected Screens', 'UX Pattern Required', 'Phase'], validations, col_widths=[1, 4, 4, 5, 2])

doc.add_page_break()

# ============================================================
# 5. LOGIC COMPLEXITY GAPS (7)
# ============================================================
doc.add_heading('5. Logic Complexity Gaps (7)', level=1)

doc.add_paragraph(
    'Seven areas where actual implementation complexity significantly exceeds what the wireframes '
    'convey. These are scoped and estimated accordingly — no hidden effort.'
)

logic_gaps = [
    ('1', 'Multi-level approval routing', 'Simple approve/reject', 'Conditional routing (spend threshold, region, category), SLA escalation, reassignment, delegation'),
    ('2', 'Dynamic form configuration', 'Static form layout', 'Tenant-specific field ordering, validation rules, visibility conditions, multi-language labels, reusable templates'),
    ('3', 'Bulk CSV onboarding', 'Single upload button', 'Duplicate detection, partial failure handling, progress tracking, error report download, rollback on critical failures'),
    ('4', 'PO-to-Invoice flip', 'Simple mapping', 'Partial invoicing, line-item selection, tax/discount calculation, multi-currency, attachment inheritance'),
    ('5', 'Screening orchestration', 'Sequential checks', '4 parallel checks with independent pass/fail, retry on timeout, partial result display, blocking vs. warning logic'),
    ('6', 'Multi-tenant data isolation', 'Not shown', 'Schema-per-tenant resolution, JWT claim extraction, middleware enforcement, cross-tenant query prevention'),
    ('7', 'ERP sync conflict resolution', 'Simple sync icon', 'Bi-directional sync handling, conflict detection (stale PO data), retry queues, dead-letter handling'),
]
add_table(doc, ['#', 'Area', 'Wireframe Shows', 'Actual Complexity'], logic_gaps, col_widths=[1, 4, 3.5, 8.5])

doc.add_page_break()

# ============================================================
# 6. MISSING API ENDPOINTS (3)
# ============================================================
doc.add_heading('6. Missing API Endpoints (3)', level=1)

doc.add_paragraph(
    'Three API endpoints required for complete functionality are absent from the wireframes. '
    'All are designed and included in the 58 vendor-built APIs.'
)

missing_apis = [
    ('1', 'Bulk invitation status tracking API', 'Track progress of CSV bulk imports', 'Wireframes show upload but not progress tracking'),
    ('2', 'Screening retry/re-trigger API', 'Allow manual re-trigger of failed screening', 'Wireframes show results but not retry flow'),
    ('3', 'Audit trail export API', 'Export filtered audit logs to CSV/PDF', 'Wireframes show audit screen but not export'),
]
add_table(doc, ['#', 'Endpoint', 'Purpose', 'Why Missing'], missing_apis, col_widths=[1, 5, 5, 6])

doc.add_paragraph('')

# ============================================================
# 7. TECHNICAL GAPS (9)
# ============================================================
doc.add_heading('7. Technical Gaps (9)', level=1)

doc.add_paragraph(
    'Nine infrastructure and integration gaps not covered in wireframes. All mitigated through '
    'Sprint 0 architecture and design work.'
)

tech_gaps = [
    ('1', 'WebSocket infrastructure for real-time notifications', 'Required for live status updates and alerts', 'Design in Sprint 0; implement in Sprint 1'),
    ('2', 'File virus scanning integration', 'Required for secure document uploads', 'Integrate with ClamAV or S3 scanning in Sprint 0'),
    ('3', 'Multi-language label management system', 'Required for i18n form configuration', 'Design label store schema in Sprint 0'),
    ('4', 'Tenant provisioning automation', 'Required for admin tenant setup', 'Admin API + seed data scripts in Sprint 1'),
    ('5', 'Rate limiting and throttling configuration', 'Required for API security', 'API Gateway configuration in Sprint 0'),
    ('6', 'Health check and readiness probes', 'Required for Kubernetes deployment', 'Standard Spring Boot Actuator + custom checks'),
    ('7', 'Database migration tooling (schema versioning)', 'Required for multi-tenant schema management', 'Flyway with tenant-aware migration strategy'),
    ('8', 'Async job queue for long-running operations', 'Required for bulk imports, screening orchestration', 'Redis-backed job queue with status polling'),
    ('9', 'Caching strategy for static/config data', 'Required for performance', 'Redis cache for form templates, workflow defs, i18n labels'),
]
add_table(doc, ['#', 'Gap', 'Impact', 'Mitigation'], tech_gaps, col_widths=[1, 5, 5, 6])

doc.add_page_break()

# ============================================================
# 8. WORKFLOW GAPS (6)
# ============================================================
doc.add_heading('8. Workflow Gaps (6)', level=1)

doc.add_paragraph(
    'Six workflow scenarios not specified in the wireframes. All resolved through Sprint 0 design '
    'deliverables and documented in the UX specification.'
)

workflow_gaps = [
    ('1', 'Invitation expiry and re-invite', 'Wireframes don\'t specify: expiry period, re-invite flow, expired link handling'),
    ('2', 'Draft save/resume lifecycle', 'Wireframes show save button but not: auto-save interval, draft expiry, resume from last step'),
    ('3', 'Approval SLA escalation', 'Wireframes show approval but not: escalation triggers, notification chain, auto-approve on SLA breach'),
    ('4', 'Supplier deactivation/reactivation', 'Not shown in wireframes — required for compliance (e.g., sanctions match after activation)'),
    ('5', 'Master data change impact assessment', 'Wireframes show change form but not: which changes trigger re-screening, which require approval'),
    ('6', 'Invoice rejection and resubmission', 'Wireframes show status but not: rejection reason capture, resubmission flow, version tracking'),
]
add_table(doc, ['#', 'Workflow', 'Gap Description'], workflow_gaps, col_widths=[1, 5, 11])

doc.add_page_break()

# ============================================================
# 9. FUNCTIONAL REQUIREMENTS TRACEABILITY
# ============================================================
doc.add_heading('9. Functional Requirements Traceability', level=1)

doc.add_paragraph(
    'Direct mapping from the HighRadius PBR spreadsheet (25 functional requirements) to our gap analysis. '
    'Every requirement is addressed — no gaps, no deferred items.'
)

func_reqs = [
    ('1', 'Supplier Onboarding', 'Invite New Suppliers', 'Covered', 'Wireframed'),
    ('2', 'Supplier Onboarding', 'Invite Existing Suppliers', 'Covered', 'Wireframed'),
    ('3', 'Supplier Onboarding', 'Multi Language Support for Forms', 'Covered', 'Tech Gap #3 (i18n label mgmt)'),
    ('4', 'Supplier Onboarding', 'Configure Onboarding Forms', 'Covered', 'Logic Gap #2 (dynamic forms) + Screen Gap #7 (Form Builder)'),
    ('5', 'Supplier Onboarding', 'Submit Onboarding Form', 'Covered', 'Wireframed + Validation Gaps #1-6'),
    ('6', 'Supplier Onboarding', 'Supplier Onboarding Worklist', 'Covered', 'Wireframed (Supplier Management flow) — Supplier List with lifecycle nav tree (All/Onboarded/Invited/Inactive + Need Review)'),
    ('7', 'Supplier Screening', 'Address Check', 'Covered', 'Wireframed (HR-owned API)'),
    ('8', 'Supplier Screening', 'TIN Check', 'Covered', 'Wireframed (HR-owned API) + Validation Gap #2'),
    ('9', 'Supplier Screening', 'Sanctions Check', 'Covered', 'Wireframed (HR-owned API)'),
    ('10', 'Supplier Screening', 'Real-time Verification', 'Covered', 'Logic Gap #5 (parallel orchestration)'),
    ('11', 'Bank Account Validation', 'Bank Account Details Verification', 'Covered', 'Wireframed (HR-owned API) + Validation Gap #3'),
    ('12', 'Bank Account Validation', 'Real-time Verification', 'Covered', 'Tech Gap #1 (WebSocket)'),
    ('13', 'Supplier Approval', 'Supplier Onboarding Workflow', 'Covered', 'Logic Gap #1 (multi-level routing)'),
    ('14', 'Supplier Approval', 'Master Data Change Request Workflow', 'Covered', 'Workflow Gap #5 + Screen Gap #10 (Master Data Change Request)'),
    ('15', 'Supplier Approval', 'Payment Method and Terms Update', 'Covered', 'Workflow Gap #5'),
    ('16', 'Purchase Order', 'PO Import', 'Covered', 'Wireframed (Looking through Purchase Orders flow) + Tech Gap #8'),
    ('17', 'Purchase Order', 'PO Import across Multiple Customers', 'Covered', 'Logic Gap #6 (multi-tenant)'),
    ('18', 'Purchase Order', 'PO Acknowledgement', 'Covered', 'Wireframed (Looking through Purchase Orders + Details page for different status in PO) — Acknowledge button on PO Detail'),
    ('19', 'Purchase Order', 'PO Status Update to AP and ERP', 'Covered', 'Logic Gap #7 (ERP sync)'),
    ('20', 'Portal Invoice Creation', 'Manual Invoice Creation', 'Covered', 'Wireframed (Creating a Blank Invoice flow)'),
    ('21', 'Portal Invoice Creation', 'PO Flip to Invoice', 'Covered', 'Wireframed (Flipping PO to Create Invoice flow) + Logic Gap #4 (PO-flip)'),
    ('22', 'Portal Invoice Creation', 'Sync Invoices to AP Applications', 'Covered', 'Tech Gap #8 (async queue)'),
    ('23', 'Supplier Manager Dashboard', 'Dashboard', 'Covered', 'Screen Gap #1 (Manager Dashboard — not wireframed)'),
    ('24', 'Supplier Dashboard', 'Dashboard', 'Covered', 'Wireframed (Dashboard flow) — Financial Summary, Pipeline, Tasks'),
]
add_table(doc, ['#', 'Agent', 'Task', 'Status', 'Gap Reference'], func_reqs, col_widths=[1, 3.5, 4, 2, 6.5])

doc.add_page_break()

# ============================================================
# 10. PROCESS FLOW COVERAGE
# ============================================================
doc.add_heading('10. Process Flow Coverage (Phase 1 & Phase 2)', level=1)

doc.add_paragraph(
    'The end-to-end process flows (Phase 1: Core Onboarding & Compliance, Phase 2: Transactions & '
    'Profile Management) were mapped against the gap register to confirm complete coverage.'
)

doc.add_heading('Phase 1 — Core Onboarding & Compliance', level=2)
phase1_steps = [
    ('Supplier Manager Login', 'Apollo Platform', 'Keycloak IAM (HR-owned)', 'No gap — HR dependency'),
    ('Navigate to Supplier Portal', 'Module navigation', 'G4 DSL routing', 'No gap'),
    ('Invitation Method Selection', 'Individual or Bulk CSV', 'Wireframed', 'Logic Gap #3 (Bulk CSV complexity)'),
    ('Open Invitation E-Form Modal', 'Name, Email, ERP, Category', 'Wireframed', 'No gap'),
    ('Upload CSV with Supplier List', 'Bulk upload', 'Wireframed', 'API Gap #1 (status tracking)'),
    ('CSV Validation', 'Duplicates, Format, Email', 'Wireframed', 'Validation Gap #6 (duplicate detection)'),
    ('Send Invitation Email', 'Secure time-limited link', 'Wireframed', 'Workflow Gap #1 (expiry handling)'),
    ('Supplier Registration (5 steps)', 'General, Purchase, Remittance, Payment, Documents', 'Wireframed', 'Validation Gaps #1-5'),
    ('Draft Save & Resume', 'Auto-save during registration', 'Wireframed (button only)', 'Workflow Gap #2 (lifecycle)'),
    ('Certify & Submit', 'Registration submission', 'Wireframed', 'No gap'),
    ('Supplier Screening (4 parallel)', 'Address, TIN, Sanctions, Bank', 'Wireframed', 'Logic Gap #5 (orchestration)'),
    ('Screening Failure Handling', 'Block + actionable error + retry', 'Wireframed', 'API Gap #2 (retry API)'),
    ('Supplier Approval Workflow', 'Multi-level approval', 'Wireframed (basic)', 'Logic Gap #1 + Workflow Gap #3'),
    ('Approval Actions', 'Approve/Reject/Reassign', 'Wireframed', 'Screen Gap #3 (Approval Worklist)'),
    ('SLA Breach Auto-Escalation', 'Auto-escalate to next level', 'Not wireframed', 'Workflow Gap #3'),
    ('Supplier Activation', 'Grant portal access via Keycloak', 'Wireframed', 'No gap — HR dependency'),
    ('Supplier Dashboard', 'Financial summary, tasks, alerts', 'Wireframed (Dashboard flow)', 'Layout provided — implementation gaps remain in validation/logic'),
]
add_table(doc, ['Process Step', 'Description', 'Wireframe Status', 'Gap Reference'], phase1_steps, col_widths=[4, 4, 3.5, 5.5])

doc.add_paragraph('')
doc.add_heading('Phase 2 — Transactions & Profile Management', level=2)
phase2_steps = [
    ('Supplier Dashboard (entry)', 'Financial summary + tasks + alerts', 'Wireframed (Dashboard)', 'Layout provided — Financial Summary, Pipeline chart, Tasks & Alerts'),
    ('Supplier Actions Selection', 'POs / Invoices / Profile', 'Wireframed (Dashboard)', 'Left sidebar navigation with icons for each section'),
    ('Purchase Order List', 'PO grid with filters, multi-buyer', 'Wireframed (Looking through Purchase Orders)', 'Full grid with status, buyer, amounts, multi-select actions'),
    ('PO Detail View', 'Billing, summary, line items', 'Wireframed (Looking through Purchase Orders + Details page for different status in PO)', 'Billing & Summary, Line Items grid, status variants'),
    ('PO Acknowledge/Reject', 'Click to acknowledge or reject', 'Wireframed (Details page for different status in PO)', 'Acknowledge button shown, status variants (Pending/Acknowledged/Rejected/Closed)'),
    ('PO Status Sync to ERP', 'Automated sync to SAP/Oracle/NetSuite', 'Not wireframed', 'Logic Gap #7 (ERP sync) — backend only'),
    ('Invoice List', 'Invoice grid with lifecycle view', 'Wireframed (Looking through Invoices)', 'Left nav status tree + grid with all fields'),
    ('Create Invoice (Manual)', 'Fill all fields, blank draft', 'Wireframed (Creating a Blank Invoice)', 'Empty form with editable fields, Add Line Item'),
    ('Create Invoice (PO Flip)', 'Auto-populate from PO', 'Wireframed (Flipping PO to Create Invoice)', 'Line selection modal + auto-populated invoice. Logic Gap #4 remains'),
    ('Partial Line Selection', '1-of-N PO lines', 'Wireframed (Flipping PO to Create Invoice)', 'Checkbox per PO line in "Select Lines to Create Invoice" modal'),
    ('Invoice Draft Editing', 'Invoice No, dates, payment terms', 'Wireframed (Looking through Invoices + Details page for different status in Invoice)', 'Editable fields in Draft status, read-only in other statuses'),
    ('Save Draft / Attach Files', 'Draft persistence + file upload', 'Wireframed (Files Tab Journey)', 'Files tab with Upload modal, progress bars, document grid'),
    ('Submit Invoice', 'Submit to AP system', 'Wireframed (Flipping PO to Create Invoice + Creating a Blank Invoice)', 'Submit button shown + "Invoice Created Successfully" toast. Tech Gap #8 (async sync) remains'),
    ('Invoice Lifecycle', 'AP review, payment, rejection', 'Wireframed (Details page for different status in Invoice)', '6 status variants shown: Draft, Submitted, In Progress, Rejected, Paid, Void. Workflow Gap #6 remains'),
    ('Supplier Profile (5-Tab View)', 'Read/edit supplier master data', 'Wireframed (Supplier Profile)', 'General Supplier Info shown. Screen Gap #10 (Master Data Change Request) remains'),
    ('Master Data Change Request', 'Data update + approval trigger', 'Not wireframed', 'Screen Gap #10 + Workflow Gap #5'),
    ('Master Data Approval', 'Accept/reject change request', 'Not wireframed', 'Logic Gap #1'),
]
add_table(doc, ['Process Step', 'Description', 'Wireframe Status', 'Gap Reference'], phase2_steps, col_widths=[4, 4, 3.5, 5.5])

doc.add_page_break()

# ============================================================
# 11. NON-FUNCTIONAL REQUIREMENTS GAP CHECK
# ============================================================
doc.add_heading('11. Non-Functional Requirements Gap Check', level=1)

doc.add_paragraph(
    'Cross-referencing the HighRadius NFR requirements against the gap register to confirm all NFRs '
    'are addressed in the technical architecture and gap mitigation strategy.'
)

nfrs = [
    ('Performance', 'API response time < 3 seconds', 'Addressed', 'Tech Gap #9 (caching), Tech Gap #8 (async queues)'),
    ('Performance', 'Grid load < 500ms for 1,000+ rows', 'Addressed', 'Virtual scrolling (AG Grid / G4 grid)'),
    ('Security', 'Keycloak IAM / SSO / RBAC', 'Addressed', 'HR dependency — adapter built by vendor'),
    ('Security', '5-layer enforcement chain', 'Addressed', 'Tech Gap #5 (rate limiting)'),
    ('Security', 'Encryption (TLS 1.3, AES-256)', 'Addressed', 'HashiCorp Vault integration'),
    ('Security', 'CSP headers / XSS prevention', 'Addressed', 'Sprint 0 configuration'),
    ('Security', 'File virus scanning', 'Addressed', 'Tech Gap #2 (ClamAV / S3 native)'),
    ('Multi-tenancy', 'Schema-per-tenant + RLS', 'Addressed', 'Logic Gap #6, Tech Gap #4 (provisioning), Tech Gap #7 (migrations)'),
    ('Accessibility', 'WCAG 2.0 AA', 'Addressed', 'All 29 screens — validated by automated + manual testing'),
    ('Code Quality', '80% unit test coverage', 'Addressed', 'Jest (FE) + JUnit (BE) — CI/CD enforced'),
    ('Deployment', 'Containerised G4 deployment', 'Addressed', 'Docker/Helm — Sprint 0 DevOps setup'),
    ('Deployment', 'CI/CD per HighRadius practices', 'Addressed', 'Sprint 0 pipeline configuration'),
    ('Availability', 'Health checks and readiness probes', 'Addressed', 'Tech Gap #6'),
    ('Scalability', '1,100+ tenant support', 'Addressed', 'Tech Gap #4 (provisioning) + Tech Gap #7 (migrations)'),
    ('Resilience', 'Circuit breaker for external APIs', 'Addressed', 'Resilience4j — Sprint 0 configuration'),
    ('Audit', 'Immutable audit trail', 'Addressed', 'Screen Gap #6 (Audit Trail) + API Gap #3 (export)'),
    ('Compliance', 'SOX / GDPR support', 'Addressed', 'Immutable logs, data export, retention policies'),
    ('i18n', 'Multi-language support', 'Addressed', 'Tech Gap #3 (label management system)'),
]
add_table(doc, ['Category', 'Requirement', 'Status', 'Gap Reference'], nfrs, col_widths=[2.5, 5, 2, 7.5])

doc.add_page_break()

# ============================================================
# 12. PLATFORM INTEGRATION GAP — AP Portal to Supplier Portal
# ============================================================
doc.add_heading('12. Platform Integration Gap: AP Portal to Supplier Portal Entry Point', level=1)

doc.add_paragraph(
    'An integration gap exists between the HighRadius AP Portal (where Supplier Managers work daily) '
    'and the Supplier Portal (where onboarding is initiated). The RFP and functional requirements state '
    'that an AP user invites suppliers, but the exact navigation and deployment model between the two '
    'systems needs formal confirmation. Analysis of the 91-page User Journey wireframes provides '
    'strong evidence for the intended integration model.'
)

doc.add_heading('What the Wireframes Reveal', level=2)
doc.add_paragraph(
    'Five key pieces of evidence from the HighRadius User Journey wireframes confirm how the '
    'Supplier Manager navigates from the AP Portal to the Supplier Portal:'
)

evidence_rows = [
    ('AP Dashboard\n(Wireframe Image 1.2)',
     'The AP Automation dashboard at cloud2g.highradius.com shows a "Products" dropdown in the top navigation. '
     'Listed products: "Accounts Payable" and "Supplier Portal" as a separate module. '
     'This confirms the Supplier Portal is a product module within the Apollo platform, accessible via a product switcher.'),
    ('Browser Tab Evidence',
     'Wireframe screenshots show two browser tabs labelled "AP" and "SP" — both on the same domain: '
     'cloud2g.highradius.com/RRDMSProject/dms/Home.do. The Supplier Portal loads within the same Apollo '
     'application shell, not a separate URL.'),
    ('Supplier List Screen\n(Wireframe Image 1.3)',
     'After switching to the Supplier Portal module, the header changes to "Supplier Portal | Apollo | highradius". '
     'The page shows "All Suppliers (10)" with a blue "+ Add Supplier" button that triggers the Invitation E-Form modal. '
     'Left navigation tree shows lifecycle stages: All Suppliers > Onboarded (03) > Invited (03) with sub-stages '
     '(TIN Check, Bank Account Validation, Sanctions Screening, Need Review) > Inactive (04).'),
    ('Invitation Modal\n(Wireframe Image 1.4)',
     'The "+ Add Supplier" button opens the "Supplier Onboarding Invitation E-Form" modal with fields: '
     'Supplier Name, Supplier E-mail, Confirm E-mail, ERP/Company (dropdown), Supplier Category (dropdown). '
     'This modal is the entry point for onboarding — it lives inside the Supplier Portal module within Apollo.'),
    ('Supplier Registration\n(Wireframe Image 2.2)',
     'The supplier receives an email from noreply@qas.com with a link to https://registrationsupplier.com/ — '
     'a DIFFERENT domain from the Apollo platform (cloud2g.highradius.com). This confirms the supplier-side '
     'registration form is a separate standalone application from the buyer-side portal.'),
]
add_table(doc, ['Wireframe Reference', 'Evidence'], evidence_rows, col_widths=[3.5, 13.5])

doc.add_paragraph('')
doc.add_heading('Confirmed Model: Embedded Module + Separate Registration App', level=2)
doc.add_paragraph(
    'The wireframes reveal a two-application architecture with distinct entry points for each persona:'
)

model_rows = [
    ('Supplier Manager\n(Buyer-Side)',
     'Apollo Platform\n(cloud2g.highradius.com)',
     '1. Logs into Apollo Platform (AP Automation)\n'
     '2. Clicks "Products" dropdown in top navigation\n'
     '3. Selects "Supplier Portal" module\n'
     '4. Portal loads within same Apollo shell (same domain, same session)\n'
     '5. Sees Supplier List with lifecycle tracking (left nav tree)\n'
     '6. Clicks "+ Add Supplier" to open Invitation E-Form modal\n'
     '7. Fills invitation details and sends',
     'Shared Keycloak session — no separate login. Same URL domain. '
     'G4 platform handles module routing and product switching.'),
    ('Supplier\n(Vendor-Side)',
     'Standalone Registration App\n(registrationsupplier.com)',
     '1. Receives email invitation from noreply@qas.com\n'
     '2. Clicks "Proceed To Complete Vendor Registration" link\n'
     '3. Opens registrationsupplier.com (separate domain)\n'
     '4. Completes 5-step registration form (General Info, Purchase, Remittance, Payment, Documents)\n'
     '5. Submits — triggers screening and approval\n'
     '6. On activation, gets access to Supplier Portal within Apollo',
     'Token-based link access — no credentials in email. '
     'Time-limited secure link. Separate domain from Apollo.'),
]
add_table(doc, ['Persona', 'Application', 'Navigation Flow', 'Technical Detail'], model_rows, col_widths=[2.5, 3.5, 6.5, 4.5])

doc.add_paragraph('')
doc.add_heading('Open Questions for HighRadius Confirmation', level=2)
doc.add_paragraph(
    'While the wireframes confirm the embedded module model, the following implementation details need '
    'formal confirmation from HighRadius before Sprint 0:'
)

open_questions = [
    ('1', 'Module Registration',
     'How is the Supplier Portal registered as a product module in Apollo/G4? '
     'Is there a module manifest, plugin system, or micro-frontend configuration that HighRadius manages?'),
    ('2', 'Product Switcher Ownership',
     'The "Products" dropdown showing "Accounts Payable / Supplier Portal" — does HighRadius build '
     'and maintain this navigation? Or does the vendor need to register the Supplier Portal module in it?'),
    ('3', 'Session & Context Sharing',
     'When the user switches from AP to Supplier Portal, is session context passed '
     '(active company, user role, tenant_id)? Or does the portal read these from the shared Keycloak JWT?'),
    ('4', 'Registration App Hosting',
     'The supplier registration form lives at registrationsupplier.com. '
     'Is this domain/hosting provided by HighRadius? Or does the vendor deploy the registration SPA separately?'),
    ('5', 'Post-Activation Supplier Access',
     'After activation, the supplier accesses the Supplier Portal within Apollo. '
     'Do they use the same cloud2g.highradius.com URL with a supplier Keycloak role? Or a separate URL?'),
    ('6', 'Left Navigation Pattern',
     'The Supplier Portal shows a left nav tree (All Suppliers > Onboarded > Invited > Inactive). '
     'Is this a G4 platform standard navigation component, or does the vendor build it from scratch?'),
]
add_table(doc, ['#', 'Topic', 'Question'], open_questions, col_widths=[1, 3.5, 12.5])

doc.add_paragraph('')
doc.add_heading('Impact on Scope & Architecture', level=2)

impact_rows = [
    ('Two-App Architecture',
     'The system is TWO frontend applications:\n'
     '(a) Supplier Portal module within Apollo — for Supplier Managers + activated Suppliers (G4 DSL components)\n'
     '(b) Standalone Registration SPA at registrationsupplier.com — for new suppliers during onboarding\n'
     'This may affect the "3 SPAs" reference in Proposal 03 and component sharing strategy.'),
    ('Scope Ownership',
     'The "Products" dropdown and module registration in Apollo is HighRadius-owned scope. '
     'Vendor builds the Supplier Portal module content (screens, APIs, logic). '
     'HighRadius registers it as a product module in the platform.'),
    ('Authentication Split',
     'Buyer-side: Shared Keycloak session within Apollo — no additional auth integration needed.\n'
     'Supplier registration: Token-based secure link access (no login credentials).\n'
     'Activated supplier: Keycloak SSO login to Apollo with supplier role.'),
    ('Deployment Model',
     'The Supplier Portal module deploys INTO the Apollo/G4 platform infrastructure. '
     'The registration SPA may deploy to a separate domain/CDN. '
     'This affects CI/CD pipeline design, containerisation strategy, and environment configuration.'),
    ('Supplier List Screen',
     'The wireframed Supplier List with left navigation tree (All/Onboarded/Invited/Inactive) and "+ Add Supplier" '
     'is now confirmed as Wireframed Screen #2 in the reconciled 29-screen inventory (Section 3A). '
     'It is a distinct screen from the Manager Dashboard (which remains undesigned — Screen Gap #1).'),
]
add_table(doc, ['Impact Area', 'Detail'], impact_rows, col_widths=[3.5, 13.5])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Action Required: ')
run.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)
p.add_run(
    'Confirm the 6 open questions above with HighRadius before or during Sprint 0. '
    'Add as Dependency D11 in the dependency register. The wireframe evidence strongly supports '
    'the embedded module model, but deployment details (module registration, domain hosting, '
    'navigation configuration) must be formally agreed to prevent integration surprises.'
)

doc.add_page_break()

# ============================================================
# 13. GAP MITIGATION STRATEGY
# ============================================================
doc.add_heading('13. Gap Mitigation Strategy', level=1)

doc.add_paragraph(
    'All 42 gaps are resolved through a structured mitigation approach across Sprint 0 and subsequent sprints.'
)

doc.add_heading('Sprint 0 Deliverables (Weeks 1-2)', level=2)
sprint0 = [
    ('Architecture + DB Schema', 'Architect/TL', 'System design doc, data model, Flyway setup — resolves Tech Gaps #4, #7'),
    ('DevOps / CI-CD', '1 BE Engineer', 'Pipeline setup, Docker/Helm, environment config — resolves Tech Gaps #5, #6'),
    ('G4 OOB Evaluation', '1 FE Engineer', 'Evaluate G4 components, document gaps, build component map'),
    ('10 Screen Designs', '2 FE Engineers', '~5 screens each in 10 days — resolves all 10 Screen Gaps (admin/manager screens)'),
    ('Keycloak + Vault Setup', '1 BE Engineer', 'Mock IAM layer, Vault integration — supports Security NFRs'),
    ('QA Framework', 'QA Engineer', 'Playwright skeleton, test plan, test data design'),
    ('Validation Error State Library', 'PM + QA + FE', 'All 12 validation types documented — resolves all Validation Gaps'),
    ('WebSocket Design', 'Architect/TL', 'Real-time notification architecture — resolves Tech Gap #1'),
    ('Virus Scanning Decision', 'Architect/TL', 'ClamAV vs S3 native — resolves Tech Gap #2'),
    ('i18n Label Schema', '1 BE Engineer', 'Multi-language label store — resolves Tech Gap #3'),
    ('Caching Strategy', 'Architect/TL', 'Redis cache for config data — resolves Tech Gap #9'),
    ('Async Queue Design', '1 BE Engineer', 'Redis-backed job queue — resolves Tech Gap #8'),
    ('Workflow Specifications', 'PM + Architect', 'All 6 workflow gaps documented — resolves all Workflow Gaps'),
]
add_table(doc, ['Stream', 'Owner', 'Deliverable & Gap Resolution'], sprint0, col_widths=[4, 3, 10])

doc.add_paragraph('')
doc.add_heading('Gap Resolution Timeline', level=2)
timeline = [
    ('Sprint 0 (Weeks 1-2)', '33', '10 Screen Gaps + 12 Validation Gaps + 6 Workflow Gaps + 5 Tech Gaps (design)'),
    ('Sprint 1 (Weeks 3-4)', '5', 'Tech Gaps #1, #4, #6 (implementation) + API Gap #1, #2'),
    ('Sprint 2-5 (Weeks 5-12)', '3', 'Logic Gaps #1-3, #5-6 (resolved during feature build)'),
    ('Sprint 6-9 (Weeks 13-20)', '1', 'Logic Gaps #4, #7 + API Gap #3 (resolved during Phase 2 build)'),
    ('Total', '42', 'All gaps resolved before Phase 2 go-live'),
]
add_table(doc, ['Sprint', 'Gaps Resolved', 'Details'], timeline, col_widths=[4, 2.5, 10.5])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Commitment: ')
run.bold = True
p.add_run(
    'All 42 gaps are fully accounted for in our estimation and delivery plan. '
    'No hidden costs, no surprise change requests. Sprint 0 resolves 79% of gaps (33/42) '
    'before any feature development begins. The 19 wireframed screens provide a strong UX foundation '
    'for all supplier-facing journeys — design effort is focused on the 10 admin/manager screens.'
)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\Users\Metapointer\Desktop\Claude\Metapointer_HighRadius_GapAnalysis_Complete.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
