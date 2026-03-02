#!/usr/bin/env python3
"""
Convert Supplier_Portal_Proposal_v3.2.md → Supplier_Portal_Proposal_v3.3.docx
Parses the markdown and generates a professionally styled Word document.
"""
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)

DARK_BLUE = RGBColor(0, 51, 102)
GREY = RGBColor(100, 100, 100)


def add_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bold_paragraph(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def italic_paragraph(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def add_table(headers, rows, style_name='Light Grid Accent 1'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style_name
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = str(cell_text)
            for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def bullet(text, level=0):
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style_name)
    add_rich_text(p, text)


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    add_rich_text(p, text)


def section_heading(number, title):
    doc.add_heading(f'{number}. {title}', level=1)


def sub_heading(text):
    doc.add_heading(text, level=2)


def sub_sub_heading(text):
    doc.add_heading(text, level=3)


def classification_note(text):
    p = doc.add_paragraph()
    r = p.add_run(f'Classification: {text}')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.italic = True


def standard_placeholder():
    italic_paragraph('[Standard section — to be populated with Metapointer standard template]')


def add_rich_text(paragraph, text):
    """Parse inline markdown bold **text** and add runs to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def add_rich_paragraph(text):
    """Add a paragraph with inline bold support."""
    p = doc.add_paragraph()
    add_rich_text(p, text)
    return p


# ============================================================
# Read and parse the markdown
# ============================================================
with open('Supplier_Portal_Proposal_v3.2.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

lines = md_content.split('\n')


def parse_table(start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = []
    rows = []
    i = start_idx

    # Parse header row
    header_line = lines[i].strip()
    if header_line.startswith('|'):
        headers = [c.strip() for c in header_line.split('|')[1:-1]]
    i += 1

    # Skip separator row
    if i < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i].strip()):
        i += 1

    # Parse data rows
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
        # Pad row if needed
        while len(row) < len(headers):
            row.append('')
        rows.append(row[:len(headers)])
        i += 1

    return headers, rows, i


def strip_bold_markers(text):
    """Remove ** markers for plain text contexts."""
    return text.replace('**', '')


# ============================================================
# COVER / TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_title('AP SUPPLIER PORTAL')
add_title('ENTERPRISE PROPOSAL')
doc.add_paragraph()
add_subtitle('HighRadius Corporation')
doc.add_paragraph()

add_table(
    ['Field', 'Detail'],
    [
        ('Version', '3.3'),
        ('Date', 'February 28, 2026'),
        ('Client', 'HighRadius Corporation'),
        ('Project', 'AP Supplier Portal Development & Deployment'),
        ('Service Provider', 'Metapointer'),
        ('Status', 'READY FOR SUBMISSION'),
    ]
)

doc.add_page_break()

# ============================================================
# Process the markdown content
# ============================================================
i = 0
# Skip past the title and metadata table
while i < len(lines):
    if lines[i].strip() == '## Table of Contents':
        break
    i += 1

# TOC
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph()

toc_parts = [
    ('PART I — PROPOSAL CONTEXT', [
        '1. Cover Page',
        '2. Document Control',
        '3. Executive Summary',
        '4. Business Understanding',
    ]),
    ('PART II — SOLUTION DESIGN', [
        '5. Scope of Work',
        '6. Solution Overview',
        '7. Solution Architecture',
        '8. Non-Functional Requirements',
    ]),
    ('PART III — ENGINEERING & DELIVERY', [
        '9. DevOps & Release Management',
        '10. Delivery Approach',
        '11. Project Plan & Milestones',
        '12. Team Structure',
    ]),
    ('PART IV — QUALITY & GOVERNANCE', [
        '13. Quality Assurance Framework',
        '14. Data Migration Strategy',
        '15. Risk Management',
        '16. Change Management Process',
    ]),
    ('PART V — SECURITY & SUPPORT', [
        '17. Security & Compliance Controls',
        '18. Documentation & Knowledge Transfer',
        '19. Support & Warranty Model',
    ]),
    ('PART VI — COMMERCIAL & LEGAL', [
        '20. Commercial Proposal',
        '21. Legal & Contractual Terms',
        '22. About Metapointer',
    ]),
]

for part_name, sections in toc_parts:
    p = doc.add_paragraph()
    r = p.add_run(part_name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_BLUE
    for section in sections:
        p = doc.add_paragraph(section)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# Skip to the actual content (after TOC in markdown)
while i < len(lines):
    if lines[i].strip().startswith('## 1. Cover Page'):
        break
    i += 1

# Section number to title mapping for page breaks
section_numbers_for_page_break = set()
for num in range(1, 23):
    section_numbers_for_page_break.add(num)

current_section_num = 0

while i < len(lines):
    line = lines[i].strip()

    # Skip empty lines
    if not line:
        i += 1
        continue

    # Horizontal rule (---)
    if line == '---':
        i += 1
        continue

    # H2 = Section heading (## N. Title)
    match_h2 = re.match(r'^## (\d+)\.\s+(.+)$', line)
    if match_h2:
        sec_num = int(match_h2.group(1))
        sec_title = match_h2.group(2)

        # Page break before each major section (except first one if already at top)
        if current_section_num > 0:
            doc.add_page_break()
        current_section_num = sec_num

        section_heading(sec_num, sec_title.upper())

        # Check next line for classification
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # Skip empty lines to find classification
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                cls_match = re.match(r'^\*\*Classification:\s*(.+?)\*\*$', lines[j].strip())
                if cls_match:
                    classification_note(cls_match.group(1))
                    i = j + 1
                    continue

        i += 1
        continue

    # Classification line
    cls_match = re.match(r'^\*\*Classification:\s*(.+?)\*\*$', line)
    if cls_match:
        classification_note(cls_match.group(1))
        i += 1
        continue

    # H3 = Sub heading (### N.N Title)
    match_h3 = re.match(r'^### (.+)$', line)
    if match_h3:
        sub_heading(match_h3.group(1))
        i += 1
        continue

    # H4 = Sub-sub heading (#### Title)
    match_h4 = re.match(r'^#### (.+)$', line)
    if match_h4:
        sub_sub_heading(match_h4.group(1))
        i += 1
        continue

    # Standard placeholder
    if line.startswith('*[Standard section') or line.startswith('*[Standard Metapointer'):
        standard_placeholder()
        i += 1
        continue

    # Table
    if line.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
        headers, rows, end_idx = parse_table(i)
        # Clean bold markers from table cells
        clean_headers = [strip_bold_markers(h) for h in headers]
        clean_rows = [[strip_bold_markers(c) for c in row] for row in rows]
        add_table(clean_headers, clean_rows)
        i = end_idx
        continue

    # Numbered list (1. **text**)
    match_num = re.match(r'^(\d+)\.\s+(.+)$', line)
    if match_num and not re.match(r'^\d+\.\s+(Cover|Document|Executive|Business|Scope|Solution|Non|DevOps|Delivery|Project|Team|Quality|Data|Risk|Change|Security|Documentation|Support|Commercial|Legal|About)', line):
        numbered(strip_bold_markers(match_num.group(2)))
        i += 1
        continue

    # Bullet points (- **text** or - text)
    match_bullet = re.match(r'^-\s+(.+)$', line)
    if match_bullet:
        bullet_text = match_bullet.group(1)
        bullet(bullet_text)
        i += 1
        continue

    # Sub-bullet (  - text)
    match_sub_bullet = re.match(r'^\s+-\s+(.+)$', line)
    if match_sub_bullet:
        bullet(match_sub_bullet.group(1), level=1)
        i += 1
        continue

    # Bold paragraph (**text:**)
    if line.startswith('**') and line.endswith('**'):
        bold_paragraph(strip_bold_markers(line))
        i += 1
        continue

    # Bold paragraph with content after
    match_bold_start = re.match(r'^\*\*(.+?)\*\*\s*$', line)
    if match_bold_start:
        bold_paragraph(match_bold_start.group(1))
        i += 1
        continue

    # Italic line
    if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
        text = line.strip('*').strip()
        italic_paragraph(text)
        i += 1
        continue

    # Regular paragraph (with potential inline bold)
    if line and not line.startswith('#') and not line.startswith('|'):
        add_rich_paragraph(line)
        i += 1
        continue

    i += 1

# ============================================================
# Footer
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Prepared By: ')
r.bold = True
p.add_run('Metapointer')

p = doc.add_paragraph()
r = p.add_run('For: ')
r.bold = True
p.add_run('HighRadius Corporation')

p = doc.add_paragraph()
r = p.add_run('Date: ')
r.bold = True
p.add_run('February 28, 2026')

# Save
output_path = 'Supplier_Portal_Proposal_v3.3.docx'
doc.save(output_path)
print(f'Successfully generated: {output_path}')
