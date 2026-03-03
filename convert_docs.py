#!/usr/bin/env python3
"""
Generic Markdown → DOCX converter for the 6 split proposal documents.
Usage: python convert_docs.py                    # converts all 6
       python convert_docs.py Doc1_Proposal_PDF  # converts one
"""
import sys
import re
import os

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DARK_BLUE = RGBColor(0, 51, 102)
GREY = RGBColor(100, 100, 100)

# All 6 document files
ALL_DOCS = [
    'Doc1_Proposal_PDF',
    'Doc2_Scope_of_Work',
    'Doc3_Timeline',
    'Doc4_Tech_Specs',
    'Doc5_Case_Studies',
    'Doc6_Contract_MSA',
    'Doc7_QA_Strategy',
]


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = DARK_BLUE


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def strip_bold_markers(text):
    return text.replace('**', '')


def add_rich_text(paragraph, text):
    """Parse inline markdown bold **text** and italic *text* and add runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def add_rich_paragraph(doc, text):
    p = doc.add_paragraph()
    add_rich_text(p, text)
    return p


def add_table(doc, headers, rows, style_name='Light Grid Accent 1'):
    if not headers or not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style_name
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = strip_bold_markers(h)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = strip_bold_markers(str(cell_text))
            for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def bullet(doc, text, level=0):
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style_name)
    add_rich_text(p, text)


def numbered_item(doc, text):
    p = doc.add_paragraph(style='List Number')
    add_rich_text(p, text)


def classification_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f'Classification: {text}')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.italic = True


def italic_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def bold_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    headers = []
    rows = []
    i = start_idx

    header_line = lines[i].strip()
    if header_line.startswith('|'):
        headers = [c.strip() for c in header_line.split('|')[1:-1]]
    i += 1

    if i < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i].strip()):
        i += 1

    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
        while len(row) < len(headers):
            row.append('')
        rows.append(row[:len(headers)])
        i += 1

    return headers, rows, i


def convert_markdown_to_docx(md_path, docx_path):
    """Convert a markdown file to a styled .docx document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)

    # --- COVER PAGE ---
    # Extract H1 title
    title_text = ''
    meta_table_start = -1
    for idx, line in enumerate(lines):
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
        if line.strip().startswith('| Field'):
            meta_table_start = idx
            break

    # Title
    doc.add_paragraph()
    if '—' in title_text:
        parts = title_text.split('—', 1)
        add_title(doc, parts[0].strip())
        add_title(doc, parts[1].strip())
    else:
        add_title(doc, title_text)

    doc.add_paragraph()
    add_subtitle(doc, 'HighRadius Corporation')
    doc.add_paragraph()

    # Metadata table
    if meta_table_start >= 0:
        headers, rows, end_idx = parse_table(lines, meta_table_start)
        add_table(doc, headers, rows)
    else:
        end_idx = 0

    doc.add_page_break()

    # --- Find where content starts (after TOC) ---
    i = end_idx
    toc_start = -1
    content_start = -1

    while i < len(lines):
        line = lines[i].strip()
        if line == '## Table of Contents':
            toc_start = i
        # First ## section after TOC
        if toc_start >= 0 and i > toc_start and re.match(r'^## \d+\.', line):
            content_start = i
            break
        i += 1

    # --- TOC ---
    if toc_start >= 0:
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph()

        j = toc_start + 1
        while j < len(lines) and (content_start < 0 or j < content_start):
            line = lines[j].strip()
            if line == '---':
                j += 1
                continue
            if not line:
                j += 1
                continue
            if line.startswith('## '):
                break

            # Parse TOC entries
            # Top level: "1. [Title](#link)"
            toc_top = re.match(r'^(\d+)\.\s+\[(.+?)\]\(.*?\)$', line)
            if toc_top:
                p = doc.add_paragraph(f'{toc_top.group(1)}. {toc_top.group(2)}')
                p.paragraph_format.space_after = Pt(2)
                j += 1
                continue

            # Sub level: "   - 1.1 [Title](#link)"
            toc_sub = re.match(r'^\s+-\s+(\d+\.\d+)\s+\[(.+?)\]\(.*?\)$', line)
            if toc_sub:
                p = doc.add_paragraph(f'{toc_sub.group(1)} {toc_sub.group(2)}')
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                j += 1
                continue

            # Plain TOC line (indented items without links)
            toc_plain = re.match(r'^\s+-\s+(.+)$', line)
            if toc_plain:
                p = doc.add_paragraph(f'  {toc_plain.group(1)}')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(0)
                j += 1
                continue

            j += 1

        doc.add_page_break()

    # --- MAIN CONTENT ---
    if content_start < 0:
        content_start = end_idx

    i = content_start
    current_section = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Horizontal rule
        if line == '---':
            i += 1
            continue

        # Footer line at end
        if line.startswith('*Prepared By:'):
            i += 1
            continue

        # H2 = Section heading (## N. Title)
        match_h2 = re.match(r'^## (\d+)\.\s+(.+)$', line)
        if match_h2:
            sec_num = int(match_h2.group(1))
            sec_title = match_h2.group(2)

            if current_section > 0:
                doc.add_page_break()
            current_section = sec_num

            doc.add_heading(f'{sec_num}. {sec_title.upper()}', level=1)

            # Check for classification line
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                cls_match = re.match(r'^\*\*Classification:\s*(.+?)\*\*$', lines[j].strip())
                if cls_match:
                    classification_note(doc, cls_match.group(1))
                    i = j + 1
                    continue

            i += 1
            continue

        # Classification line (standalone)
        cls_match = re.match(r'^\*\*Classification:\s*(.+?)\*\*$', line)
        if cls_match:
            classification_note(doc, cls_match.group(1))
            i += 1
            continue

        # H3 = Sub heading
        match_h3 = re.match(r'^### (.+)$', line)
        if match_h3:
            doc.add_heading(match_h3.group(1), level=2)
            i += 1
            continue

        # H4 = Sub-sub heading
        match_h4 = re.match(r'^#### (.+)$', line)
        if match_h4:
            doc.add_heading(match_h4.group(1), level=3)
            i += 1
            continue

        # Standard placeholder
        if line.startswith('*[Standard section') or line.startswith('*[Standard Metapointer') or line.startswith('*[To be populated'):
            text = line.strip('*').strip()
            italic_paragraph(doc, text)
            i += 1
            continue

        # Blockquote (> text) — render as italic indented
        if line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            if quote_text.startswith('**') and ':**' in quote_text:
                # Bold label in blockquote
                label_match = re.match(r'^\*\*(.+?)\*\*\s*(.*)', quote_text)
                if label_match:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    r = p.add_run(label_match.group(1))
                    r.bold = True
                    r.font.size = Pt(10)
                    if label_match.group(2):
                        r2 = p.add_run(' ' + label_match.group(2))
                        r2.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    add_rich_text(p, quote_text)
            elif quote_text.startswith('-'):
                # Bullet inside blockquote
                bullet_text = quote_text.lstrip('-').strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)
                add_rich_text(p, f'• {bullet_text}')
                for run in p.runs:
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                add_rich_text(p, quote_text)
                for run in p.runs:
                    run.font.size = Pt(10)
            i += 1
            continue

        # Table
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i + 1].strip()):
            headers, rows, end_idx = parse_table(lines, i)
            add_table(doc, headers, rows)
            i = end_idx
            continue

        # Numbered list
        match_num = re.match(r'^(\d+)\.\s+(.+)$', line)
        if match_num:
            # Don't treat TOC-like lines as numbered items
            text = match_num.group(2)
            if not re.match(r'\[.+\]\(.+\)', text):
                numbered_item(doc, text)
                i += 1
                continue

        # Bullet points (- text)
        match_bullet = re.match(r'^-\s+(.+)$', line)
        if match_bullet:
            bullet(doc, match_bullet.group(1))
            i += 1
            continue

        # Sub-bullet (  - text)
        match_sub_bullet = re.match(r'^\s+-\s+(.+)$', line)
        if match_sub_bullet:
            bullet(doc, match_sub_bullet.group(1), level=1)
            i += 1
            continue

        # Bold-only paragraph
        if line.startswith('**') and line.endswith('**'):
            bold_paragraph(doc, strip_bold_markers(line))
            i += 1
            continue

        # Italic-only line
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            text = line.strip('*').strip()
            italic_paragraph(doc, text)
            i += 1
            continue

        # Regular paragraph
        if line and not line.startswith('#'):
            add_rich_paragraph(doc, line)
            i += 1
            continue

        i += 1

    # --- FOOTER ---
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Prepared By: ')
    r.bold = True
    p.add_run('Metapointer')

    p = doc.add_paragraph()
    r = p.add_run('For: ')
    r.bold = True
    p.add_run('HighRadius Corporation')

    p = doc.add_paragraph()
    r = p.add_run('Date: ')
    r.bold = True
    p.add_run('February 28, 2026')

    doc.save(docx_path)
    return docx_path


def main():
    args = sys.argv[1:]

    if args:
        # Convert specific docs
        targets = []
        for arg in args:
            name = arg.replace('.md', '').replace('.docx', '')
            if name in ALL_DOCS:
                targets.append(name)
            else:
                print(f'Unknown document: {arg}. Available: {", ".join(ALL_DOCS)}')
                sys.exit(1)
    else:
        targets = ALL_DOCS

    for name in targets:
        md_path = f'{name}.md'
        docx_path = f'{name}.docx'

        if not os.path.exists(md_path):
            print(f'SKIP: {md_path} not found')
            continue

        print(f'Converting {md_path} → {docx_path} ...', end=' ')
        try:
            convert_markdown_to_docx(md_path, docx_path)
            print('OK')
        except Exception as e:
            print(f'FAILED: {e}')

    print('\nDone.')


if __name__ == '__main__':
    main()
