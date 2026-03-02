#!/usr/bin/env python3
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)

DARK_BLUE = RGBColor(0, 51, 102)
GREY = RGBColor(100, 100, 100)


def add_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bold_paragraph(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def italic_paragraph(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def add_table(headers, rows, style_name='Light Grid Accent 1'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style_name
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = str(cell_text)
            for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def bullet(text, level=0):
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2'
    doc.add_paragraph(text, style=style_name)


def numbered(text):
    doc.add_paragraph(text, style='List Number')


def section_heading(number, title):
    doc.add_heading(f'{number}. {title}', level=1)


def sub_heading(text):
    doc.add_heading(text, level=2)


def sub_sub_heading(text):
    doc.add_heading(text, level=3)


def standard_placeholder():
    italic_paragraph('[Standard section \u2014 to be populated with Metapointer standard template]')


def classification_note(text):
    p = doc.add_paragraph()
    r = p.add_run(f'Classification: {text}')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.italic = True


# ============================================================
# COVER / TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_title('AP SUPPLIER PORTAL')
add_title('ENTERPRISE PROPOSAL')
doc.add_paragraph()
add_subtitle('HighRadius Corporation')
doc.add_paragraph()

add_table(
    ['Field', 'Detail'],
    [
        ('Version', '3.2'),
        ('Date', 'February 28, 2026'),
        ('Client', 'HighRadius Corporation'),
        ('Project', 'AP Supplier Portal Development & Deployment'),
        ('Service Provider', 'Metapointer'),
        ('Status', 'READY FOR SUBMISSION'),
    ]
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph()

toc_parts = [
    ('PART I — PROPOSAL CONTEXT', [
        '1. Cover Page',
        '2. Document Control',
        '3. Executive Summary',
        '4. Business Understanding',
    ]),
    ('PART II — SOLUTION DESIGN', [
        '5. Scope of Work',
        '6. Solution Overview',
        '7. Solution Architecture',
        '8. Non-Functional Requirements',
    ]),
    ('PART III — ENGINEERING & DELIVERY', [
        '9. DevOps & Release Management',
        '10. Delivery Approach',
        '11. Project Plan & Milestones',
        '12. Team Structure',
    ]),
    ('PART IV — QUALITY & GOVERNANCE', [
        '13. Quality Assurance Framework',
        '14. Data Migration Strategy',
        '15. Risk Management',
        '16. Change Management Process',
    ]),
    ('PART V — SECURITY & SUPPORT', [
        '17. Security & Compliance Controls',
        '18. Documentation & Knowledge Transfer',
        '19. Support & Warranty Model',
    ]),
    ('PART VI — COMMERCIAL & LEGAL', [
        '20. Commercial Proposal',
        '21. Legal & Contractual Terms',
        '22. About Metapointer',
    ]),
]

for part_name, sections in toc_parts:
    p = doc.add_paragraph()
    r = p.add_run(part_name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_BLUE
    for section in sections:
        p = doc.add_paragraph(section)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ============================================================
# 1. COVER PAGE
# ============================================================
section_heading(1, 'COVER PAGE')
classification_note('PROJECT-SPECIFIC')

bullet('Client: HighRadius Corporation')
bullet('Project: AP Supplier Portal \u2014 Development & Deployment')
bullet('Proposal Version: 3.2')
bullet('Date: February 28, 2026')
bullet('Prepared By: Metapointer')

doc.add_page_break()

# ============================================================
# 2. DOCUMENT CONTROL
# ============================================================
section_heading(2, 'DOCUMENT CONTROL')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# 3. EXECUTIVE SUMMARY
# ============================================================
section_heading(3, 'EXECUTIVE SUMMARY')
classification_note('PROJECT-SPECIFIC')

sub_heading('3.2 Business Context')
doc.add_paragraph(
    'HighRadius is a fintech enterprise SaaS provider offering an Agentic AI platform for the Office of the CFO. '
    'Its cloud solutions span the end-to-end Order-to-Cash (O2C) cycle, Accounts Payable, Treasury Management, '
    'Record-to-Report, and B2B Payments \u2014 integrating 180+ AI agents (microservices) to orchestrate financial '
    'processes. Adopted by over 1,100 leading companies (including 3M, Unilever, Procter & Gamble, and '
    'Johnson & Johnson), HighRadius is consistently recognized as a market leader by Gartner, IDC, and Forrester.'
)
doc.add_paragraph(
    'As part of its expanding AP automation capabilities, HighRadius requires a robust Supplier Portal '
    'to streamline vendor management operations across its customer base. The portal will serve as a '
    'unified platform for buyers to onboard suppliers, manage vendor relationships, and for suppliers '
    'to get onboarded, receive POs, submit and track invoices. This aligns with HighRadius\u2019s AI-first '
    'platform strategy and its mission to automate financial processes in multi-tenant environments handling '
    'diverse ERPs like SAP, Oracle, and NetSuite. By integrating with HighRadius\u2019s G4 platform, the portal '
    'will enhance AP efficiency, support acquisitions, and provide analytics-driven insights.'
)

sub_heading('3.2 Problem Statement')
doc.add_paragraph(
    'Current AP operations at HighRadius rely on manual, fragmented processes that create significant '
    'operational and financial risk:'
)

bullet('Onboarding Delays: Email-based supplier onboarding takes 15\u201320 days with inconsistent data collection, processing 50\u2013100+ suppliers per quarter with extended activation delays, lost early-payment discounts, and strained buyer-supplier relationships')
bullet('Compliance Gaps: Only ~30% of suppliers are screened, creating exposure to sanctions violations ($20M+ OFAC fines), tax/identity fraud from error-prone TIN verification, and banking vulnerabilities from unvalidated changes \u2014 with no centralized audit trails')
bullet('Invoice Processing: Manual channels (email/fax) yield only ~60% SLA compliance with high error rates, late penalties, and forecasting inaccuracies due to 12\u201315 manual touchpoints per event across hundreds to thousands of invoices monthly')
bullet('Operational Inefficiency: 60\u201370% of AP time spent on transactional work rather than strategic activities; 100% reliance on buyer intervention for every supplier interaction')
bullet('Visibility & Data Silos: Supplier data scattered across ERPs, spreadsheets, and emails \u2014 no single source of truth, manual reports, and delayed decision-making')

sub_heading('3.3 Proposed Solution')
doc.add_paragraph(
    'The Supplier Portal is structured around 7 intelligent agents (4 automated, 3 assisted) across '
    'three personas, delivering 29 screens + 11 modals with 66 APIs (58 vendor-built, 7 integrations, '
    '1 orchestration) supporting 25+ core features:'
)

add_table(
    ['Agent Name', 'Type', 'Key Capabilities'],
    [
        ('Supplier Onboarding', 'Automated', 'End-to-end invitation, dynamic multi-language form configuration, document upload/validation, draft save/resume'),
        ('Supplier Screening', 'Automated', 'Real-time compliance verification \u2014 address, TIN, and sanctions checks in parallel'),
        ('Bank Account Validation', 'Automated', 'Financial security through real-time bank details verification and ownership validation'),
        ('Supplier Approval', 'Assisted', 'Multi-level approval workflows with configurable routing, SLA-based escalation'),
        ('Master Data Change Approval', 'Assisted', 'Orchestrates approval workflows for master data updates with conditional routing'),
        ('Purchase Order', 'Automated', 'ERP ingestion (multi-customer), acknowledgment/reject synchronization'),
        ('Portal Invoice Creation', 'Assisted', 'Manual entry and PO-flip capability with validation, attachments, and AP synchronization'),
    ]
)

bold_paragraph('Technology Stack:')
bullet('Frontend: React 18.x + HighRadius G4 DSL with WCAG 2.0 AA accessibility')
bullet('Backend: Spring Boot 3.x (Java 17+), microservices architecture')
bullet('Database: PostgreSQL OLAP with multi-tenant 100% data isolation (schema-per-tenant, row-level security)')
bullet('Security: Keycloak IAM integration, HashiCorp Vault secrets management, TLS 1.3, AES-256 encryption')
bullet('Deployment: Docker containers, Kubernetes orchestration in HighRadius G4 environment')

doc.add_paragraph()
bold_paragraph('AI-Powered Productivity:')
doc.add_paragraph(
    'Leveraging Claude Opus AI for code generation, testing automation, and documentation, '
    'delivering 1.10x\u20131.50x productivity gains with 6 core engineers.'
)

sub_heading('3.4 Timeline Summary')
bullet('Kickoff: March 2026')
bullet('Sprint 0 (Foundation): 2 weeks \u2014 architecture, DevOps setup, G4 component evaluation & scaffolding, design sprint')
bullet('Phase 1 (MVP): 4 agents (Onboarding, Screening, Bank Validation, Approval) \u2014 Sprints 1\u20135 (10 weeks) \u2014 UAT + 1 live customer by mid-May 2026')
bullet('Phase 2: 3 agents (Purchase Order, Portal Invoice Creation, Master Data Change Approval) \u2014 Sprints 6\u20139 (8 weeks) \u2014 Full delivery by early July 2026')
bullet('Hypercare: 6 months post-go-live through January 2027 with production support and knowledge transfer')
bullet('Total Duration: 20 calendar weeks aligned to RFP milestones with built-in risk buffers')

sub_heading('3.5 Commercial Summary')
bullet('Engagement Model: Fixed-Price with defined scope and deliverables')
bullet('Payment Structure: Milestone-based with Go/No-Go gates (5 milestones: 20%\u201330%\u201320%\u201320%\u201310%)')
bullet('Detailed Pricing: Provided in separate Pricing Agreement document')
bullet('Change Management: Formal change control process for scope modifications')

doc.add_page_break()

# ============================================================
# 4. BUSINESS UNDERSTANDING
# ============================================================
section_heading(4, 'BUSINESS UNDERSTANDING')
classification_note('PROJECT-SPECIFIC')

sub_heading('4.1 Current State')
doc.add_paragraph(
    'HighRadius\u2019s current accounts payable (AP) operations are characterized by manual and '
    'fragmented supplier management processes, lacking a dedicated self-service portal:'
)
bullet('Onboarding: Primarily email-driven, involving unstructured workflows with inconsistent data collection and extended cycles of 15\u201320 days')
bullet('Invoice Processing: Relies on disparate channels (email, fax, manual entry), leading to ~60% SLA compliance and elevated error rates')
bullet('Compliance: Reactive and ad-hoc, with screening covering only ~30% of suppliers and no centralized audit trails')
bullet('Supplier Interactions: All interactions require full buyer intervention with no self-service capabilities for PO acknowledgment or invoice tracking')
bullet('Data Management: Scattered across ERPs (SAP, Oracle, NetSuite), spreadsheets, and emails \u2014 poor real-time visibility and manual reporting')

doc.add_paragraph(
    'This setup supports three main personas \u2014 Supplier Manager (handling invitations and workflows), '
    'Supplier (submitting documents manually), and HighRadius Admin/Client Admin (managing configurations) '
    '\u2014 but in an inefficient, non-digital manner.'
)

sub_heading('4.2 Key Challenges')
bullet('Manual Onboarding Delays: Processing 50\u2013100+ suppliers per quarter via email leads to activation delays, lost early-payment discounts, and strained buyer-supplier relationships')
bullet('Invoice and PO Processing Inefficiencies: Handling hundreds to thousands of invoices monthly through fragmented channels causes late penalties, forecasting inaccuracies, and high manual touchpoints (12\u201315 per event), diverting AP resources from strategic tasks')
bullet('Compliance and Risk Gaps: With ~70% of suppliers unscreened, there is significant exposure to regulatory violations (OFAC sanctions fines up to $20M+), tax/identity fraud from error-prone TIN verification, and banking fraud due to unvalidated changes \u2014 exacerbated by scattered audit trails')
bullet('Lack of Self-Service and Scalability: 100% reliance on buyer intervention hampers scaling, especially for multi-tenant environments or acquisitions, leading to high per-supplier costs and inability to absorb growth without proportional headcount increases')
bullet('Data Fragmentation: Siloed data across 3+ systems results in no single source of truth, manual reports, delayed decisions, and opportunity costs (60\u201370% of AP time on transactional work)')

sub_heading('4.3 Target State')
doc.add_paragraph('The Supplier Portal will transform HighRadius\u2019s AP operations into a fully digitized, automated, and scalable ecosystem:')

add_table(
    ['Area', 'Current State', 'Target State'],
    [
        ('Supplier Onboarding', '15\u201320 days, email-driven', '3\u20135 days, automated with real-time screening'),
        ('Screening Coverage', '~30% of suppliers screened', '100% automated compliance checks'),
        ('Invoice SLA Compliance', '~60%', '>90% with automated validation'),
        ('First-Pass Accuracy', '~65%', '>95% via PO-flip and field validation'),
        ('Supplier Self-Service', '0% (full buyer intervention)', '>80% adoption target'),
        ('Manual Touchpoints', '12\u201315 per event', '80%+ reduction through automation'),
        ('Audit Trail', 'Scattered across emails/spreadsheets', 'Centralized, immutable, 2-year retention'),
        ('Data Visibility', 'Manual reports, delayed decisions', 'Real-time dashboards and KPI tracking'),
    ]
)

sub_heading('4.4 Business Outcomes & ROI Indicators')
doc.add_paragraph('The Supplier Portal is expected to deliver measurable business value across multiple dimensions:')

add_table(
    ['Outcome Area', 'Metric', 'Expected Impact'],
    [
        ('Cost Savings', 'FTE equivalent saved through automation', '2\u20133 FTEs redirected from transactional to strategic work'),
        ('Late Payment Penalties', 'Reduction in missed SLA penalties', 'Significant reduction through >90% SLA compliance'),
        ('Compliance Risk', 'Non-compliant supplier exposure', 'Eliminated \u2014 100% screening coverage with full audit trail'),
        ('Supplier Satisfaction', 'Self-service adoption & onboarding speed', '>80% adoption; 3\u20135 day onboarding vs 15\u201320 days'),
        ('Scalability', 'Ability to absorb growth without headcount', 'Multi-tenant automation handles acquisitions without linear staffing'),
        ('Data-Driven Decisions', 'Real-time visibility', 'KPI dashboards replace manual reports; proactive management'),
    ]
)

bold_paragraph('Value Realization Timeline:')
bullet('Months 1\u20133 (Post Go-Live): Onboarding cycle reduction, screening automation, initial self-service adoption')
bullet('Months 3\u20136: Invoice SLA improvements, dashboard-driven decisions, supplier satisfaction gains')
bullet('Months 6\u201312: Full ROI realization \u2014 cost savings, compliance maturity, scalability proven through tenant expansion')

sub_heading('4.5 Assumptions')
bullet('HighRadius will provide timely access to owned services and dependencies: Keycloak IAM, screening/bank validation APIs, ERP samples (SAP/Oracle/NetSuite), G4 platform for deployment, and test data/environments for UAT')
bullet('The RFP scope (7 agents, specified NFRs) remains fixed; changes managed through formal change control with impact assessments')
bullet('Weekly UX/UI sign-offs and stakeholder feedback will be provided by HighRadius to avoid design iteration delays')
bullet('Multi-tenancy requirements align with HighRadius standards for 100% data isolation with tenant resolution via JWT claims')
bullet('Project kickoff assumes March 2026 start, with no major external disruptions')
bullet('Performance testing will use production-scale datasets provided by HighRadius')
bullet('Product requirements and UX specifications will be provided by HighRadius as per RFP Section 5')

sub_heading('4.6 Constraints')
bullet('Platform Constraint: All deployment must target HighRadius G4 containerized environment \u2014 no standalone infrastructure or alternative cloud hosting')
bullet('Technology Constraint: Frontend must use React 18.x with G4 DSL components; backend must use Spring Boot 3.x (Java 17+) per HighRadius engineering standards')
bullet('Timeline Constraint: March 2026 kickoff is non-negotiable; any delay in kickoff shifts all milestones proportionally')
bullet('Team Size Constraint: Fixed team of 7 FTE (6 engineers + 1 PM) \u2014 scope adjustments required if team availability changes')
bullet('Scope Freeze Constraint: RFP-defined scope (7 agents, 29 screens, specified NFRs) is frozen post-Sprint 0; changes require formal CR process with timeline/cost impact assessment')
bullet('Integration Constraint: Vendor has no control over HighRadius-owned services (Keycloak, screening APIs, bank validation, G4 platform) \u2014 delivery timelines depend on these services being available and stable')
bullet('Compliance Constraint: All code must pass HighRadius infosec clearance before production deployment; no exceptions for timeline acceleration')

doc.add_page_break()

# ============================================================
# 5. SCOPE OF WORK
# ============================================================
section_heading(5, 'SCOPE OF WORK')
classification_note('PROJECT-SPECIFIC')

sub_heading('5.1 In Scope')

sub_sub_heading('Application Development & Quantified Deliverables')
bold_paragraph('Supplier Portal Web Application (responsive, modern UI for three personas):')
bullet('29 application screens across Supplier Manager, Supplier, and HighRadius Admin/Client Admin personas', level=1)
bullet('11 popup/modals for workflows and confirmations', level=1)
bullet('66 APIs: 58 vendor-built, 7 integrations, 1 orchestration layer', level=1)
bullet('25+ core features including: invitation-based onboarding, dynamic multi-language forms, document upload/validation, draft save/resume, real-time screening (TIN/address/sanctions/bank), configurable approval workflows, SLA-based escalation, PO import/sync, invoice creation, master data change workflows, manager dashboards', level=1)
bullet('Responsive design with WCAG 2.0 AA accessibility standards', level=1)
bullet('Virtual scrolling for grids with 1,000+ rows', level=1)
bullet('Integration of 91-page user journey wireframes', level=1)

sub_sub_heading('Screen & Workflow Estimation')
bold_paragraph('Portal Screens Breakdown (29 Total):')

add_table(
    ['Workflow/Feature', 'Screens', 'Complexity', 'Est. Hours/Screen', 'Total Hours'],
    [
        ('Supplier Onboarding', '6', 'High', '40', '240'),
        ('Supplier Screening', '3', 'Medium', '35', '105'),
        ('Bank Validation', '2', 'Medium', '30', '60'),
        ('Approval Workflows', '4', 'High', '45', '180'),
        ('PO Management', '5', 'High', '40', '200'),
        ('Invoice Management', '5', 'High', '40', '200'),
        ('Master Data Management', '2', 'Medium', '35', '70'),
        ('Dashboards', '2', 'High', '50', '100'),
    ]
)

doc.add_paragraph('Modals & Pop-ups (11 Total): 15\u201320 hours each = ~200 hours')
bold_paragraph('Screen Development Total: ~1,355 hours')

doc.add_paragraph()
bold_paragraph('Workflow/Feature Development (Backend Services):')

add_table(
    ['Agent/Service', 'Workflows', 'Est. Hours/Workflow', 'Total Hours'],
    [
        ('Supplier Onboarding Agent', '3', '120', '360'),
        ('Supplier Screening Agent', '2', '100', '200'),
        ('Bank Validation Agent', '1', '80', '80'),
        ('Supplier Approval Agent', '2', '140', '280'),
        ('Master Data Approval Agent', '2', '120', '240'),
        ('Purchase Order Agent', '3', '130', '390'),
        ('Portal Invoice Agent', '3', '130', '390'),
    ]
)

bold_paragraph('Agent/Service Development Total: ~1,940 hours')

sub_sub_heading('Integrations')
bullet('Keycloak IAM for OIDC/SSO with RBAC/ABAC')
bullet('Screening APIs (address, TIN, sanctions) \u2014 parallel checks')
bullet('Bank validation API')
bullet('ERP/AP for PO/invoice sync (SAP/Oracle/NetSuite)')
bullet('Document storage (S3-compatible with virus scans)')
bullet('HashiCorp Vault for secrets management')
bullet('Notification services (email/WebSocket)')
bullet('Fault tolerance with circuit breakers, retries, and dead-letter queues')

sub_sub_heading('Testing & Quality Assurance')

bold_paragraph('Unit Testing:')
add_table(
    ['Area', 'Target Coverage', 'Estimated Effort'],
    [
        ('Frontend Unit Tests (React/Jest)', '80% coverage (RFP mandate)', '~5 weeks (QA)'),
        ('Backend Unit Tests (JUnit)', '80% coverage (RFP mandate)', '~6 weeks (QA)'),
        ('Regression Test Suite (JUnit)', 'Automated regression per RFP requirement', '~3 weeks (QA)'),
    ]
)

bold_paragraph('Integration & E2E Testing:')
add_table(
    ['Area', 'Scope', 'Effort'],
    [
        ('Supplier Onboarding', 'Invite \u2192 Form fill \u2192 Screening \u2192 Approval \u2192 Activation (6 features)', '~4 weeks (QA)'),
        ('Screening & Bank Validation (Integration)', 'Verify vendor orchestration calls HR screening/bank APIs correctly; validate result display and error handling', '~2 weeks (QA)'),
        ('Approval Workflows', 'Multi-level onboarding approval, Master data change, Payment terms (3 features)', '~3 weeks (QA)'),
        ('PO / Invoice Flows', 'PO import \u2192 Acknowledge \u2192 Flip to Invoice \u2192 Submit \u2192 AP Sync (7 features)', '~5 weeks (QA)'),
        ('Document Management', 'Upload, Download, Delete, Version control, Bulk operations', '~1.5 weeks (QA)'),
        ('Dashboard & Reporting', 'KPI calculations, Data aggregation, Drill-downs', '~1.5 weeks (QA)'),
        ('Portal Cross-cutting Flows', 'Keycloak SSO token flow \u2192 RBAC \u2192 Multi-tenant isolation \u2192 Session mgmt', '~2 weeks (QA)'),
    ]
)

bold_paragraph('Performance & Load Testing:')
add_table(
    ['Test Area', 'Benchmark', 'Effort'],
    [
        ('API Response Time Validation', 'All APIs < 3 seconds (RFP requirement)', '~1.5 weeks (QA)'),
        ('Grid Load Time', '< 500ms for 1000 rows (NFR requirement)', '~1 week (QA)'),
        ('Virtual Scrolling Stress', '> 1000 rows, 30+ columns performance', '~1 week (QA)'),
        ('Concurrent User Load', 'Simulate 100\u2013500 concurrent suppliers', '~1.5 weeks (QA)'),
        ('ERP Sync Throughput', 'PO import + Invoice sync under load', '~1 week (QA)'),
        ('Horizontal Scaling Validation', 'Container scaling under G4 environment', '~1 week (QA)'),
    ]
)

bold_paragraph('Security Testing:')
add_table(
    ['Security Area', 'Scope', 'Effort'],
    [
        ('Vulnerability Scanning (OWASP Top 10)', 'SAST + DAST scanning of all vendor-built endpoints', '~1.5 weeks (QA)'),
        ('SSO Token Validation Testing', 'Keycloak token consumption, session handling, expiry edge cases', '~1 week (QA)'),
        ('RBAC / Permission Boundary Testing', 'Role-based access enforcement across all 29 screens + 11 popups', '~1.5 weeks (QA)'),
        ('Data Protection & Isolation', 'Multi-tenant data isolation verification, PII handling', '~1 week (QA)'),
        ('CSP & Headers Validation', 'Content Security Policy, unsafe-eval prevention', '~3 days (QA)'),
        ('HashiCorp Vault Integration Test', 'Key store operations, secret rotation', '~3 days (QA)'),
        ('Infosec Clearance Prep', 'Documentation and remediation for HighRadius cybersecurity team', '~1 week (QA)'),
    ]
)

bold_paragraph('ABAC / RBAC Testing:')
add_table(
    ['ABAC/RBAC Area', 'Scope', 'Effort'],
    [
        ('Role Matrix Validation', 'Supplier, Supplier Manager, Admin, Approver \u2014 all 29 screens + 11 popups', '~1.5 weeks (QA)'),
        ('Access Boundary Enforcement', 'Cross-tenant data access prevention, Row-level security', '~1 week (QA)'),
        ('Permission Escalation Testing', 'Ensure no privilege escalation via API manipulation', '~1 week (QA)'),
        ('Multi-Tenant RBAC Isolation', 'Roles scoped to tenant, no cross-tenant role leakage', '~1 week (QA)'),
    ]
)

sub_sub_heading('Post-Deployment Support')
bullet('6-month hypercare at 0.5 FTE')
bullet('Production monitoring and issue resolution per SLAs')
bullet('Weekly updates (email/call/PVA)')
bullet('Root cause analysis and knowledge transfer')

sub_sub_heading('AI Productivity Enhancements')
doc.add_paragraph('All 6 engineers use Claude Opus AI subscriptions. Productivity multipliers applied:')

add_table(
    ['Activity', 'AI Multiplier'],
    [
        ('Code Generation (Frontend)', '1.40x'),
        ('Code Generation (Backend)', '1.30x'),
        ('Unit Test Writing', '1.40x'),
        ('API Documentation', '1.50x'),
        ('Integration Code', '1.15x'),
        ('DevOps & Config', '1.10x'),
    ]
)

doc.add_paragraph('These are conservative multipliers assuming experienced developers using AI as an accelerator, not a replacement for engineering judgment.')

sub_sub_heading('Responsibility Matrix (RACI)')
doc.add_paragraph('Clear ownership delineation between Metapointer (Vendor) and HighRadius \u2014 critical for fixed-price scoping:')

add_table(
    ['Capability', 'Vendor', 'HighRadius', 'Notes'],
    [
        ('UI/UX Development (29 screens + 11 modals)', 'R, A', 'C, I', 'Vendor builds end-to-end; HR provides wireframes & sign-off'),
        ('Backend Services (58 vendor-built APIs)', 'R, A', 'C', 'Vendor full ownership'),
        ('Orchestration Layer (1 API)', 'R, A', 'C', 'Vendor builds orchestration calling HR services'),
        ('Keycloak IAM (Auth/Login/SSO)', 'I', 'R, A', 'HighRadius-owned; vendor integrates only'),
        ('Screening Services (Address/TIN/Sanctions)', 'C', 'R, A', 'HighRadius-owned; vendor integrates + stores results'),
        ('Bank Validation Service', 'C', 'R, A', 'HighRadius-owned; vendor integrates only'),
        ('ERP Integration (SAP/Oracle/NetSuite)', 'R', 'A, C', 'Vendor builds adapters; HR provides API access & samples'),
        ('Document Storage (S3-compatible)', 'R', 'C', 'Vendor builds integration; HR provides infrastructure'),
        ('G4 Platform / Deployment Environment', 'C', 'R, A', 'HighRadius provides; vendor deploys containers'),
        ('Multi-Tenant Architecture', 'R, A', 'C', 'Vendor designs & implements isolation'),
        ('Testing & QA', 'R, A', 'C', 'Vendor owns; HR provides test data & UAT sign-off'),
        ('DevOps / CI-CD Pipeline', 'R', 'A, C', 'Vendor sets up per HR engineering practices'),
        ('Product Requirements & UX Specs', 'C', 'R, A', 'HighRadius provides; vendor implements'),
        ('Hypercare Support (6 months)', 'R', 'A, I', 'Vendor provides 0.5 FTE'),
        ('Knowledge Transfer & Handover', 'R', 'A', 'Work completion requires HR engineering sign-off'),
    ]
)
doc.add_paragraph('R = Responsible (does the work) | A = Accountable (final sign-off) | C = Consulted | I = Informed')

sub_sub_heading('UI Gap Analysis Summary')
doc.add_paragraph('Analysis of the 91-page HighRadius User Journey wireframes revealed significant gaps that are accounted for in our estimation:')

add_table(
    ['Gap Category', 'Count', 'Details'],
    [
        ('Screens Provided by HighRadius', '13', 'From User Journey wireframes'),
        ('Missing / To-Be-Designed Screens', '16', '55% of screens require UX design from scratch'),
        ('Validation Gaps', '12', 'Validation types specified but not illustrated in wireframes'),
        ('Logic Complexity Gaps', '7', 'Areas with more implementation complexity than wireframes suggest'),
        ('Missing API Endpoints', '3', 'Required endpoints not reflected in wireframes'),
        ('Technical Gaps', '9', 'Infrastructure/integration gaps not covered in wireframes'),
        ('Workflow Gaps', '6', 'Workflow logic not fully specified in wireframes'),
        ('Total Gaps', '48', 'All accounted for in estimation and delivery plan'),
    ]
)

doc.add_paragraph(
    'Mitigation: Validation Error State Library to be created in Sprint 0. All error messages, toast notifications, '
    'and inline validation indicators documented before Sprint 1 development begins. Sprint 0 design sprint addresses '
    'missing screens, with UX running 1 sprint ahead of development throughout.'
)

sub_sub_heading('Frontend Component Inventory')

add_table(
    ['Category', 'Count', 'Examples'],
    [
        ('G4 OOB Components (Reused)', '~15\u201320', 'Grids, form controls, buttons, modals, navigation \u2014 leveraged directly from G4 DSL library'),
        ('Custom Components (G4-Extended)', '~10', 'Portal-specific components not covered by G4 OOB (e.g., screening progress, approval chain visualization, PO-flip mapper)'),
        ('Screen-Level React Components', '39', 'Across 19 screens, mapping directly to dev tasks and Jira tickets'),
        ('Layout Definitions', '8', 'Page templates for different screen types'),
        ('Total Components', '~72\u201377', 'G4-first approach reduces custom build effort'),
    ]
)

doc.add_paragraph('G4 OOB component evaluation completed in Sprint 0. Custom components built only where G4 library gaps exist. Component-per-screen mapping drives sprint planning and task allocation.')

sub_heading('5.2 Out of Scope')
bullet('HighRadius-Owned Services Build: Development or modification of Keycloak authentication/login/signup, bank validation service, screening services (address/TIN/sanctions), or G4 OOB components \u2014 vendor scope limited to integration effort only')
bullet('Additional ERP Integrations: ERP types beyond SAP/Oracle/NetSuite or custom multi-customer mappings exceeding RFP specs')
bullet('Data Migration: Migration of existing supplier data from ERPs, spreadsheets, or other legacy systems')
bullet('Infrastructure Provisioning: Hardware, cloud resources, or environment setup beyond containerized G4 deployment')
bullet('Testing Exclusions: Real-world production data testing (vendor uses mocks/samples); compliance certifications beyond SOX/GDPR support')
bullet('Scope Changes: Requirements exceeding defined scope without approved change requests, including major wireframe revisions after design sign-off')

doc.add_page_break()

# ============================================================
# 6. SOLUTION OVERVIEW
# ============================================================
section_heading(6, 'SOLUTION OVERVIEW')
classification_note('PROJECT-SPECIFIC')

sub_heading('6.1 Solution Vision')
doc.add_paragraph(
    'The Supplier Portal is a secure, scalable, self-service web application designed to digitize and '
    'streamline supplier onboarding, compliance validation, purchase order management, and invoice '
    'processing for HighRadius\u2019s AP automation platform.'
)

bold_paragraph('The solution aims to:')
bullet('Eliminate manual, email-driven supplier interactions by providing a single system of engagement')
bullet('Enforce compliance and validation at the point of data entry through automated real-time checks')
bullet('Improve operational efficiency through automation, reducing onboarding cycles from 15\u201320 days to 3\u20135 days and invoice SLAs from ~60% to >90%')
bullet('Provide real-time visibility to suppliers and internal stakeholders via dashboards')
bullet('Support multi-tenant, enterprise-grade scalability with 100% data isolation')

sub_heading('6.2 Capability Map')

add_table(
    ['Agent', 'Type', 'Key Features'],
    [
        ('Supplier Onboarding', 'Automated', 'Individual/bulk (CSV) invitations, dynamic multi-language forms, document upload/validation, draft save/resume, duplicate checking'),
        ('Supplier Screening', 'Automated', 'Parallel real-time checks: address verification, TIN validation, sanctions screening (OFAC), with progress indicators and actionable error feedback'),
        ('Bank Account Validation', 'Automated', 'Real-time bank detail verification, ownership validation, fraud prevention'),
        ('Supplier Approval', 'Assisted', 'Multi-level configurable workflows, conditional routing (e.g., spend threshold), SLA auto-escalation, approve/reject/reassign'),
        ('Master Data Change Approval', 'Assisted', 'Approval workflows for supplier data updates, payment terms changes, configurable chains'),
        ('Purchase Order', 'Automated', 'Multi-ERP PO import/sync, acknowledgment/reject by supplier, real-time status tracking'),
        ('Portal Invoice Creation', 'Assisted', 'Manual entry + one-click PO-flip, partial invoicing, tax/discount lines, attachments, AP sync'),
    ]
)

bold_paragraph('Additional Self-Service Capabilities:')
bullet('Supplier Dashboard: invoice status pipeline, aging summary, payment status, tasks & alerts')
bullet('Manager Dashboard: onboarding funnel, pending tasks, SLA compliance, KPI tracking (pass/fail rates, coverage %)')
bullet('In-app/email notifications and document management with versioning')
bullet('Notification Center with real-time alerts for PO receipts, invoice status changes, and approval decisions')

sub_heading('6.3 High-Level Workflow')
numbered('Supplier Manager sends onboarding invitation (individual or bulk via CSV)')
numbered('Supplier completes self-registration: multi-step dynamic form, document upload, real-time validation, with draft save/resume and multi-language support')
numbered('Automated compliance screening triggered: address, TIN, sanctions, and bank validation run in parallel with progress indicators')
numbered('Submission enters approval workflow: multi-level, configurable routing with SLA-based escalation')
numbered('Upon approval, supplier is activated, notified, and granted full portal access via Keycloak SSO')
numbered('Supplier views/acknowledges/rejects POs (imported from ERPs), creates/submits invoices (manual or PO-flip)')
numbered('Invoice and PO data syncs with ERP/AP systems with pre-submission field-level validation')
numbered('Stakeholders track progress via real-time dashboards, notification center, and exportable reports')

sub_heading('6.4 RFP Requirement Traceability')
doc.add_paragraph('Direct mapping from RFP requirements to proposed solution modules, confirming full coverage:')

add_table(
    ['RFP Requirement Area', 'Solution Module', 'Screens', 'Phase', 'Coverage'],
    [
        ('Supplier Onboarding (invitations, forms, documents)', 'Supplier Onboarding Agent', '6 + 3 modals', 'Phase 1', 'Full'),
        ('Compliance Screening (address, TIN, sanctions)', 'Supplier Screening Agent', '3 + 1 modal', 'Phase 1', 'Full'),
        ('Bank Account Validation', 'Bank Account Validation Agent', '2', 'Phase 1', 'Full'),
        ('Multi-Level Approval Workflows', 'Supplier Approval Agent', '4 + 2 modals', 'Phase 1', 'Full'),
        ('Master Data Change Management', 'Master Data Change Approval Agent', '2 + 1 modal', 'Phase 2', 'Full'),
        ('PO Import/Acknowledgment (multi-ERP)', 'Purchase Order Agent', '5 + 2 modals', 'Phase 2', 'Full'),
        ('Invoice Creation (manual + PO-flip)', 'Portal Invoice Creation Agent', '5 + 2 modals', 'Phase 2', 'Full'),
        ('Dashboards (Supplier + Manager)', 'Cross-cutting', '2', 'Phase 1', 'Full'),
        ('Keycloak IAM / SSO / RBAC', 'Integration (HR-owned)', 'Login/Auth', 'Phase 1', 'Integration only'),
        ('Multi-Tenancy & Data Isolation', 'Architecture (schema-per-tenant)', 'All screens', 'Phase 1', 'Full'),
        ('WCAG 2.0 AA Accessibility', 'Frontend architecture', 'All screens', 'Both', 'Full'),
        ('80% Code Coverage', 'QA framework', '\u2014', 'Both', 'Full'),
        ('<3s API Response', 'Backend architecture', '\u2014', 'Both', 'Full'),
        ('Containerized G4 Deployment', 'DevOps/CI-CD', '\u2014', 'Phase 1', 'Full'),
        ('Hypercare (6 months)', 'Support model', '\u2014', 'Post-deploy', 'Full'),
        ('Handover (GIT, docs, KT)', 'Delivery governance', '\u2014', 'Phase 2 end', 'Full'),
    ]
)

bold_paragraph('Result: 100% RFP requirement coverage across all 7 agents, NFRs, engagement expectations, and success criteria.')

sub_heading('6.5 Persona Journeys')

bold_paragraph('Supplier Journey:')
doc.add_paragraph(
    'Receives email invitation with secure, time-limited link \u2192 completes multi-step registration \u2192 '
    'automated screening \u2192 approval workflow \u2192 dashboard access \u2192 PO management \u2192 invoice creation '
    'via PO-flip \u2192 status tracking'
)

bold_paragraph('Supplier Manager Journey:')
doc.add_paragraph(
    'SSO login \u2192 Manager Dashboard with pipeline metrics \u2192 send invitations (individual/bulk) \u2192 '
    'monitor worklist with status filters \u2192 review screening results \u2192 process approval tasks \u2192 '
    'track SLA compliance'
)

bold_paragraph('Admin Journey:')
doc.add_paragraph(
    'Configure onboarding forms (drag-and-drop builder, multi-language labels, reusable templates) \u2192 '
    'define approval workflows (chains, escalation rules, conditional routing) \u2192 manage users/roles '
    'via Keycloak \u2192 view audit trail (filterable, exportable)'
)

doc.add_page_break()

# ============================================================
# 7. SOLUTION ARCHITECTURE
# ============================================================
section_heading(7, 'SOLUTION ARCHITECTURE')
classification_note('PROJECT-SPECIFIC')

sub_heading('7.1 Architecture Overview')
doc.add_paragraph(
    'The architecture follows a modular, layered approach designed to meet HighRadius NFRs: sub-3-second '
    'API responses, 100% multi-tenant data isolation, containerized G4 deployment, and Keycloak-based '
    'authentication. Clear separation between presentation, API gateway, service, integration, and data '
    'layers ensures maintainability and independent scalability.'
)

sub_heading('7.2 Application Architecture')
bullet('Presentation Layer: React 18.x + G4 DSL, responsive SPA with role-based routing, WCAG 2.0 AA compliance, ARIA tags, virtual scrolling')
bullet('API Gateway: Request routing, JWT validation, tenant context injection, rate limiting')
bullet('Service Layer: Spring Boot 3.x microservices with stateless design supporting horizontal scaling; workflow-driven business logic for onboarding, approvals, and invoice processing')
bullet('Integration Layer: Adapter pattern with circuit breaker, retry, and dead-letter queues for all external system communication')

sub_heading('7.3 Data Architecture')
bullet('Multi-Tenancy Model: Schema-per-tenant with shared infrastructure; each HighRadius customer gets an isolated PostgreSQL schema with row-level security policies')
bullet('Common Schema: Shared configuration (form templates, workflow definitions, i18n labels)')
bullet('Tenant Resolution: JWT claim (tenant_id extracted from Keycloak token) applied at middleware layer before any DB query')
bullet('Core Entities: Supplier, Invitation, Screening Result, Approval Workflow, Purchase Order, Invoice, Document, Audit Log')
bullet('Strategy: PostgreSQL OLAP with read replicas, Redis caching layer, immutable audit logs with 2-year retention')

sub_heading('7.4 Integration Architecture')

add_table(
    ['External System', 'Integration Method', 'Ownership'],
    [
        ('Keycloak IAM', 'OIDC/SSO, JWT validation, RBAC', 'HighRadius-owned; vendor integrates'),
        ('Screening Services (Address/TIN/Sanctions)', 'REST API, parallel execution', 'HighRadius-owned; vendor integrates'),
        ('Bank Validation Service', 'REST API', 'HighRadius-owned; vendor integrates'),
        ('ERP Systems (SAP/Oracle/NetSuite)', 'REST API for PO import/invoice sync', 'Vendor builds adapter/orchestration'),
        ('Document Storage', 'S3-compatible API with virus scanning', 'Vendor builds integration'),
        ('HashiCorp Vault', 'Secrets management API', 'HighRadius-provided; vendor integrates'),
        ('Notification Service', 'Email/WebSocket for real-time alerts', 'Vendor builds'),
    ]
)

doc.add_paragraph('All integrations use fault tolerance patterns: circuit breaker, exponential retry, dead-letter queues.')

sub_heading('7.5 Security Architecture')
bullet('Authentication: Keycloak SSO with OIDC/JWT; no request reaches business logic without validated token, tenant context, and role verification (zero-trust model)')
bullet('Authorization: Hybrid RBAC + ABAC \u2014 roles assigned in Keycloak (RBAC), runtime attributes for fine-grained control (ABAC)')
bullet('Enforcement Points: Frontend route guards \u2192 API Gateway JWT + RBAC middleware \u2192 Service Layer ABAC checks \u2192 Database tenant_id WHERE clause + row-level security \u2192 Audit logging of all decisions')
bullet('Encryption: TLS 1.3 in transit, AES-256 at rest')
bullet('Secrets Management: HashiCorp Vault for all credentials, API keys, certificates')
bullet('Audit: Immutable audit logs with full traceability \u2014 filterable by user, action, date; exportable for compliance')

sub_heading('7.6 Deployment Architecture')
bullet('Containerization: Docker containers with Helm charts for Kubernetes orchestration')
bullet('Environment Promotion: Dev (local) \u2192 QA (auto on merge) \u2192 UAT (manual, QA sign-off) \u2192 Production (manual, UAT + HighRadius Engineering sign-off)')
bullet('High Availability: Redundant components, health checks, auto-scaling policies')
bullet('G4 Platform: Deployed within HighRadius G4 environment per RFP requirements')

doc.add_page_break()

# ============================================================
# 8. NON-FUNCTIONAL REQUIREMENTS
# ============================================================
section_heading(8, 'NON-FUNCTIONAL REQUIREMENTS')
classification_note('PROJECT-SPECIFIC')

sub_heading('8.1 Availability')
bullet('99.5% system availability target')
bullet('Redundant components to minimize single points of failure')
bullet('Health checks and auto-restart policies')
bullet('Planned maintenance windows with minimal disruption')

sub_heading('8.2 Performance')

add_table(
    ['Metric', 'Target'],
    [
        ('API Response Time (P95)', '< 3 seconds'),
        ('API Response Time (P99)', '< 3 seconds'),
        ('Grid Load (1,000+ rows)', '< 500 milliseconds'),
        ('Page Load Time', '< 2 seconds'),
        ('Screening Execution', 'Real-time with progress indicators'),
        ('File Upload/Download', '< 5 seconds for typical documents'),
    ]
)

bullet('Asynchronous processing for long-running operations (screening, bulk imports)')
bullet('Efficient pagination and filtering for large datasets')

sub_heading('8.3 Scalability')
bullet('1,000 concurrent users / 10,000 transactions per hour')
bullet('Horizontal scalability \u2014 frontend, backend, and integration layers scale independently')
bullet('Architecture supports onboarding additional tenants without impacting existing users')
bullet('Multi-tenant automation absorbs acquisitions without proportional headcount')

sub_heading('8.4 Security')
bullet('Keycloak-based SSO with OIDC/JWT authentication')
bullet('RBAC/ABAC enforcement at every layer (zero-trust)')
bullet('TLS 1.3 for data in transit, AES-256 for data at rest')
bullet('HashiCorp Vault for secrets management')
bullet('OWASP Top 10 compliance')
bullet('Secure file handling with virus scanning for document uploads')
bullet('CSP headers and XSS protection')

sub_heading('8.5 Compliance')
bullet('Full immutable audit trail for all onboarding, approval, and invoice actions')
bullet('2-year audit log retention, filterable and exportable')
bullet('SOX/GDPR support alignment')
bullet('100% screening coverage eliminating non-compliant supplier risk')
bullet('Data retention and traceability aligned with enterprise audit standards')

sub_heading('8.6 Code Quality')
bullet('80% minimum test code coverage (backend \u2014 JUnit, frontend \u2014 Jest)')
bullet('Static code analysis with no critical/high findings')
bullet('PR review policy: every merge requires at least 1 peer review')
bullet('Definition of Done includes: code complete, tests passing, reviewed, integration tested, security scanned, documented')

doc.add_page_break()

# ============================================================
# 9. DEVOPS & RELEASE MANAGEMENT
# ============================================================
section_heading(9, 'DEVOPS & RELEASE MANAGEMENT')
classification_note('STANDARD CORE + PROJECT ADAPTATION')

italic_paragraph('[Standard Metapointer CI/CD framework applies]')

doc.add_paragraph()
bold_paragraph('Project-Specific Adaptations:')
bullet('CI/CD pipeline follows HighRadius engineering practices (per RFP Section 5)')
bullet('Deployment to HighRadius G4 containerized environment using Docker/Helm')
bullet('Environment promotion: Dev \u2192 QA (auto on merge) \u2192 UAT (manual, QA sign-off) \u2192 Production (manual, UAT + HighRadius Engineering sign-off)')
bullet('Rollback procedures defined for each deployment stage')
bullet('Pre-deployment checklist: all automated tests passing, code review completed, security scan passed, documentation updated, stakeholder approval obtained')
bullet('Adherence to HighRadius cybersecurity policies during development')

doc.add_page_break()

# ============================================================
# 10. DELIVERY APPROACH
# ============================================================
section_heading(10, 'DELIVERY APPROACH')
classification_note('STANDARD CORE + PROJECT ADAPTATION')

italic_paragraph('[Standard Metapointer Agile delivery methodology applies]')

doc.add_paragraph()
bold_paragraph('Project-Specific Adaptations:')
bullet('Engagement Model: Fixed-Price with defined scope and milestone-based payments')
bullet('Sprint Cadence: 2-week sprints (10 sprints total: Sprint 0 + Sprints 1\u20139)')
bullet('Phase 1 (MVP): Sprints 0\u20135 \u2014 4 agents (Onboarding, Screening, Bank Validation, Approval)')
bullet('Phase 2: Sprints 6\u20139 \u2014 3 agents (Purchase Order, Invoice Creation, Master Data Change Approval)')
bullet('Communication Cadence: Weekly status calls with HighRadius, weekly email updates and PVA (Product Velocity Analysis) on deliverables')
bullet('Governance: Steering committee reviews, defect management SLAs, change control per Section 16')
bullet('MASC (Success Criteria): (a) UAT delivery, (b) 1 live customer in production, (c) 90%+ KPI achievement per agent')

doc.add_paragraph()
bold_paragraph('Vendor Commitments per RFP Section 5:')
bullet('Design collaboration with HighRadius engineering team with timely sign-offs')
bullet('Feature-level technical design documentation')
bullet('Functional and regression test cases (JUnit + functional)')
bullet('QA completed before delivery to HighRadius')
bullet('Handover repository: GIT, design docs, KT sessions, all project resources')

doc.add_paragraph()
bold_paragraph('Sign-Off Gates:')
bullet('Product Management: 100% module/feature delivery')
bullet('Engineering: Architecture & handover sign-off')
bullet('Infosec: Cybersecurity clearance')

doc.add_paragraph()
bold_paragraph('Phase-Level Entry & Exit Criteria:')
add_table(
    ['Phase', 'Entry Criteria', 'Exit Criteria'],
    [
        ('Sprint 0 (Foundation)', 'Signed SOW; HighRadius API access granted; G4 environment available; UX wireframes provided', 'Architecture document signed off; CI/CD pipeline operational; G4 OOB component evaluation complete; Validation Error State Library created; DB schema v1 approved'),
        ('Phase 1 Development (Sprints 1\u20135)', 'Sprint 0 exit criteria met; design sprint outputs approved; Keycloak integration environment ready', '4 agents feature-complete; 17 screens implemented; ~40 APIs deployed to QA; unit tests at 80% coverage'),
        ('Phase 1 QA + UAT', 'All Phase 1 code merged to QA branch; test data provided by HighRadius; QA environment stable', 'E2E tests passing; API response <3s validated; Grid <500ms validated; WCAG audit passed; UAT sign-off from HighRadius; 1 live customer in production'),
        ('Phase 2 Development (Sprints 6\u20139)', 'Phase 1 in production; Phase 2 design specs approved; ERP sample data for PO/Invoice flows available', '3 agents feature-complete; 12 screens implemented; ~26 APIs deployed to QA; unit tests at 80% coverage'),
        ('Phase 2 QA + Stabilization', 'All Phase 2 code merged; full regression suite ready; performance test environment configured', 'Full regression passed; virtual scrolling stress test passed; penetration test completed; UAT feedback addressed'),
        ('Full Delivery + Handover', 'Phase 2 UAT sign-off; all defects resolved or deferred with HighRadius approval', 'Production deployment complete; GIT handover accepted; KT sessions delivered; MASC criteria (a)(b)(c) achieved; infosec clearance obtained'),
    ]
)

doc.add_page_break()

# ============================================================
# 11. PROJECT PLAN & MILESTONES
# ============================================================
section_heading(11, 'PROJECT PLAN & MILESTONES')
classification_note('PROJECT-SPECIFIC')

sub_heading('11.1 Phase Breakdown')

bold_paragraph('Sprint 0 \u2014 Foundation (Weeks 1\u20132):')
bullet('Architecture finalization, design sprint for undesigned screens')
bullet('DevOps/CI-CD pipeline setup per HighRadius practices')
bullet('G4 OOB component evaluation and custom component scaffolding')
bullet('Validation Error State Library creation')
bullet('Integration environment setup with HighRadius APIs')

bold_paragraph('Phase 1 \u2014 MVP: Core Onboarding & Compliance (Sprints 1\u20135, Weeks 3\u201312):')
bullet('Supplier Onboarding Agent (invitation, dynamic forms, document upload, draft/resume)')
bullet('Supplier Screening Agent (address, TIN, sanctions \u2014 parallel real-time checks)')
bullet('Bank Account Validation Agent (bank detail verification)')
bullet('Supplier Approval Agent (multi-level workflows, SLA escalation)')
bullet('Core dashboards (Supplier Dashboard, Manager Dashboard)')
bullet('Initial integrations (Keycloak, screening APIs, bank validation)')
bullet('UAT completion and 1 live customer in production')

bold_paragraph('Phase 2 \u2014 Transactional Enablement (Sprints 6\u20139, Weeks 13\u201320):')
bullet('Purchase Order Agent (multi-ERP import, acknowledge/reject, sync)')
bullet('Portal Invoice Creation Agent (manual entry, PO-flip, AP sync)')
bullet('Master Data Change Approval Agent (update workflows)')
bullet('Enhanced reporting and analytics')
bullet('Full ERP synchronization')
bullet('Production-ready deployment and handover')

bold_paragraph('Hypercare & Stabilization (6 months post-go-live):')
bullet('0.5 FTE dedicated support rotating across team')
bullet('Production monitoring and issue resolution per SLAs')
bullet('Performance tuning and optimization')
bullet('Knowledge transfer to HighRadius operations team')
bullet('Weekly updates via email/call + PVA')

sub_heading('11.2 Timeline')

add_table(
    ['Sprint', 'Weeks', 'Focus', 'Milestone'],
    [
        ('Sprint 0', '1\u20132', 'Foundation, architecture, DevOps setup', 'Design Sign-off'),
        ('Sprint 1\u20132', '3\u20136', 'Onboarding + Screening agents', '\u2014'),
        ('Sprint 3\u20134', '7\u201310', 'Bank Validation + Approval + Dashboards', '\u2014'),
        ('Sprint 5', '11\u201312', 'Integration testing, UAT prep', 'MVP UAT Complete'),
        ('\u2014', '~Mid-May 2026', 'MVP Go-Live', '1 Live Customer'),
        ('Sprint 6\u20137', '13\u201316', 'PO Agent + Invoice Agent', '\u2014'),
        ('Sprint 8\u20139', '17\u201320', 'Master Data Agent + Full integration', 'Phase 2 UAT Complete'),
        ('\u2014', '~Early Jul 2026', 'Full Portal Go-Live', 'Full Delivery & Handover'),
        ('\u2014', 'Jul 2026 \u2013 Jan 2027', 'Hypercare', 'Hypercare Closure'),
    ]
)

bold_paragraph('Total Program Duration: 20 calendar weeks (Sprint 0 through Handover) + 6 months hypercare')

sub_heading('11.3 Calendar Estimates by Discipline')
doc.add_paragraph('Based on actual team composition (3 FE + 2 BE + 1 QA, all with Claude Opus AI). All estimates are AI-adjusted:')

add_table(
    ['Discipline', 'Team Size', 'Working Days', 'Calendar Weeks', 'Notes'],
    [
        ('Frontend Development', '3 developers', '72 days', '~15 weeks', 'Parallel execution across 3 devs; AI multiplier on code + unit tests'),
        ('Backend Development', '2 developers', '77 days', '~16 weeks', 'Includes DevOps, observability; AI multiplier on services + APIs'),
        ('QA Engineering', '1 engineer', '~44 weeks total', '~40% overlaps with dev', 'E2E, performance, security, WCAG; unit tests written by devs'),
        ('Sprint 0 (Foundation)', 'Full team', '10 days', '2 weeks', 'Architecture, design sprint, CI/CD setup, G4 component evaluation'),
    ]
)

bullet('Frontend and backend streams run in parallel, with QA overlapping ~40% with active development sprints')
bullet('Critical path is backend at ~16 weeks, with frontend completing within the same window')
bullet('QA effort extends beyond dev completion for final regression, performance, and security testing')

sub_heading('11.4 Timeline Reality Assessment')
doc.add_paragraph('The HighRadius RFP defines three delivery milestones. Our honest assessment with a 7-person team:')

add_table(
    ['RFP Milestone', 'RFP Target', 'Our Realistic Target', 'Gap', 'Notes'],
    [
        ('MVP (4 agents)', 'Feb 2026', 'Mid-May 2026', '~2.5 months', 'MVP and Phase 1 combined (both = 4 agents, 1 month apart in RFP)'),
        ('Phase 1 (4 agents)', 'March 2026', 'Mid-May 2026', '~2 months', 'Treated as combined Phase 1/MVP delivery'),
        ('Phase 2 (3 agents)', 'June 2026', 'Early July 2026', '~2\u20133 weeks', 'Mitigatable by overlapping FE/BE work streams across phases'),
    ]
)

doc.add_paragraph(
    'Recommendation: Deliver in two phases aligned to RFP agent structure. Phase 1 (4 agents) targets production '
    'by mid-May 2026 with 1 live customer. Phase 2 (3 agents) targets full delivery by early July 2026. The '
    '2\u20133 week gap from the June RFP target is mitigatable by overlapping frontend/backend work streams. '
    'Total: 20 calendar weeks from kickoff.'
)

sub_heading('11.5 Key Deliverables')

add_table(
    ['Phase', 'Calendar Window', 'Team', 'Deliverables'],
    [
        ('Sprint 0 (Setup + Design)', 'Weeks 1\u20132', 'All 7 engineers', 'CI/CD (following HR practices), G4 onboarding + OOB component evaluation, Keycloak + Vault setup, design system tokens, DB schema v1, Validation Error State Library'),
        ('Phase 1 Development', 'Weeks 3\u20138', '3 FE + 2 BE + 1 QA', '4 agents: Onboarding, Screening, Bank Validation, Approval. 17 screens, ~40 APIs'),
        ('Phase 1 QA + UAT + Production', 'Weeks 9\u201312', '3 FE + 2 BE + 1 QA', 'E2E testing, JUnit regression, perf validation (API<3s, Grid<500ms), WCAG audit, UAT, 1st live customer'),
        ('Phase 2 Development', 'Weeks 13\u201316', '3 FE + 2 BE + 1 QA', '3 agents: PO, Invoice, Master Data Change. 12 screens, ~26 APIs'),
        ('Phase 2 QA + Stabilization', 'Weeks 17\u201318', '3 FE + 2 BE + 1 QA', 'E2E testing, full regression, virtual scrolling perf test, pen test, UAT feedback'),
        ('Full Delivery + Handover', 'Weeks 19\u201320', '3 FE + 2 BE + 1 QA', 'Phase 2 production deploy, feature-level tech docs, GIT handover, KT, MASC milestone (c)'),
    ]
)

bold_paragraph('Cross-Phase Deliverables:')
bullet('Security configuration and multi-tenant isolation')
bullet('CI/CD deployment pipelines per HighRadius practices')
bullet('Feature-level technical design documentation')
bullet('OpenAPI/Swagger API documentation')
bullet('Functional and regression test suites (JUnit + functional)')
bullet('Deployment guides and runbooks')
bullet('Knowledge transfer sessions')

doc.add_page_break()

# ============================================================
# 12. TEAM STRUCTURE
# ============================================================
section_heading(12, 'TEAM STRUCTURE')
classification_note('STANDARD CORE + PROJECT ALLOCATION')

italic_paragraph('[Standard Metapointer role descriptions and responsibilities apply]')

doc.add_paragraph()
bold_paragraph('Project-Specific Allocation:')

add_table(
    ['Role', 'Count', 'FTE-Months', 'Project Focus'],
    [
        ('Architect / Technical Lead', '1', '6', 'System design, G4 integration, code quality, performance optimization'),
        ('Senior Backend Engineers', '2', '12', 'Microservices, API implementation, database design, DevOps'),
        ('Frontend Engineers', '2', '12', 'UI/UX with G4 DSL, responsive design, WCAG compliance'),
        ('QA Engineer', '1', '6', 'E2E automation, performance testing, security validation, WCAG testing'),
        ('Project Manager', '1', '6', 'Timeline management, stakeholder communication, risk tracking, PVA reporting'),
    ]
)

bold_paragraph('Total Team Capacity: 7 FTE (6 months core development + extended hypercare phase at 0.5 FTE)')

doc.add_paragraph()
doc.add_paragraph(
    'AI Tooling: All 6 engineers equipped with Claude Opus AI subscriptions for 1.10x\u20131.50x '
    'productivity gains across coding, testing, and documentation activities.'
)

doc.add_page_break()

# ============================================================
# 13. QUALITY ASSURANCE FRAMEWORK
# ============================================================
section_heading(13, 'QUALITY ASSURANCE FRAMEWORK')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# 14. DATA MIGRATION STRATEGY
# ============================================================
section_heading(14, 'DATA MIGRATION STRATEGY')
classification_note('CONDITIONAL \u2014 PROJECT-SPECIFIC')

bold_paragraph('Not Applicable.')
doc.add_paragraph(
    'Data migration of existing supplier data from ERPs, spreadsheets, or legacy systems is explicitly '
    'out of scope for this engagement. The Supplier Portal will start with fresh data, with suppliers '
    'onboarded through the new portal workflows.'
)

doc.add_page_break()

# ============================================================
# 15. RISK MANAGEMENT
# ============================================================
section_heading(15, 'RISK MANAGEMENT')
classification_note('STANDARD FRAMEWORK + PROJECT RISKS')

italic_paragraph('[Standard Metapointer risk methodology applies]')

doc.add_paragraph()
bold_paragraph('Project-Specific Risks:')

add_table(
    ['#', 'Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ('1', 'G4 Platform Integration Delays: G4 OOB components or environment not ready on time', 'Medium', 'High', 'Early Sprint 0 integration testing; mock services as fallback; weekly sync with HighRadius platform team'),
        ('2', 'HighRadius API Instability: Screening, bank validation, or Keycloak APIs unstable or undocumented', 'Medium', 'High', 'API contract validation in Sprint 0; adapter pattern with circuit breakers; fallback error handling'),
        ('3', 'Multi-Tenancy Complexity: Data isolation errors or tenant configuration conflicts', 'Low', 'Critical', 'Schema-per-tenant design validation; automated isolation testing; row-level security enforcement'),
        ('4', 'Undesigned Screens (16 of 29): 55% of screens require UX design from scratch', 'Medium', 'Medium', 'Sprint 0 design sprint; UX running 1 sprint ahead of dev; weekly design reviews with HighRadius'),
        ('5', 'Workflow Configuration Complexity: Complex approval chains causing routing conflicts', 'Low', 'Medium', 'Clear workflow design docs; validation before activation; escalation logic testing'),
        ('6', 'Performance at Scale: Large PO/invoice datasets affecting grid performance', 'Medium', 'Medium', 'Virtual scrolling, pagination, query optimization, load testing before go-live'),
        ('7', 'ERP Integration Variability: Different ERP formats/behaviors across SAP/Oracle/NetSuite', 'Medium', 'Medium', 'Adapter pattern per ERP; early integration with sample data; dedicated Sprint 0 API mapping'),
        ('8', 'Regulatory/Compliance Changes: Updated sanctions lists or compliance requirements mid-project', 'Low', 'Medium', 'Modular validation architecture; configurable rule engine; decoupled screening service'),
        ('9', 'User Adoption Risk: Low supplier adoption of self-service portal', 'Low', 'Medium', 'Intuitive UX design, onboarding communication templates, dashboard visibility, training support'),
        ('10', 'Requirement Mismatch at UAT: Delivered features not matching requirement intent', 'Medium', 'Medium', 'Sprint 0 design sprint + UX 1 sprint ahead + weekly design reviews + written confirmation of ambiguous requirements'),
    ]
)

doc.add_page_break()

# ============================================================
# 16. CHANGE MANAGEMENT PROCESS
# ============================================================
section_heading(16, 'CHANGE MANAGEMENT PROCESS')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# 17. SECURITY & COMPLIANCE CONTROLS
# ============================================================
section_heading(17, 'SECURITY & COMPLIANCE CONTROLS')
classification_note('STANDARD CORE + PROJECT ADAPTATION')

italic_paragraph('[Standard Metapointer security controls apply]')

doc.add_paragraph()
bold_paragraph('Project-Specific Adaptations:')
bullet('Access Control: Keycloak integration with hybrid RBAC/ABAC model; 4 personas (Supplier, Supplier Manager, Approver, Admin) with screen-level and button-level permissions across 29 screens + 11 modals')
bullet('Encryption: TLS 1.3 in transit, AES-256 at rest; HashiCorp Vault for all secrets management per HighRadius security requirements')
bullet('Audit Logging: Immutable audit trail for all portal actions \u2014 onboarding, screening, approvals, invoice submissions; filterable by user, action type, date; exportable for compliance; 2-year retention')
bullet('Multi-Tenant Isolation: Schema-per-tenant with row-level security; tenant_id enforced at middleware layer before any database query; JWT-based tenant resolution')
bullet('Regulatory Alignment: SOX/GDPR support; OFAC sanctions screening compliance; 100% supplier screening coverage; OWASP Top 10 compliance')
bullet('Adherence to HighRadius cybersecurity policies during development per RFP Section 5')
bullet('Infosec clearance from HighRadius cybersecurity team as sign-off gate')

doc.add_page_break()

# ============================================================
# 18. DOCUMENTATION & KNOWLEDGE TRANSFER
# ============================================================
section_heading(18, 'DOCUMENTATION & KNOWLEDGE TRANSFER')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# 19. SUPPORT & WARRANTY MODEL
# ============================================================
section_heading(19, 'SUPPORT & WARRANTY MODEL')
classification_note('STANDARD CORE + PROJECT TERMS')

italic_paragraph('[Standard Metapointer support structure applies]')

doc.add_paragraph()
bold_paragraph('Project-Specific Terms:')
bullet('Hypercare Duration: 6 months post-go-live (approximately July 2026 \u2013 January 2027)')
bullet('Hypercare Staffing: 0.5 FTE dedicated support resource rotating across the 6-person engineering team (~13 weeks of part-time effort)')

doc.add_paragraph()
bold_paragraph('Support SLAs:')

add_table(
    ['Priority', 'Description', 'Response Time'],
    [
        ('P1', 'Critical \u2014 system down', '< 4 hours'),
        ('P2', 'High \u2014 major feature impaired', '< 8 hours'),
        ('P3', 'Medium \u2014 workaround available', '< 24 hours'),
        ('P4', 'Low \u2014 cosmetic/minor', 'Next business day'),
    ]
)

bullet('Reporting: Weekly updates via email/call and PVA (Product Velocity Analysis) on deliverables per RFP Section 5')
bullet('Escalation: Tiered escalation \u2014 Tier 1 (24h) \u2192 Tier 2 (48h) \u2192 Tier 3 with auto-escalation and incident logging')
bullet('Handover Repository: GIT repository, design documents, KT sessions, deployment runbooks, and all resources used for the product per RFP Section 6')

doc.add_page_break()

# ============================================================
# 20. COMMERCIAL PROPOSAL
# ============================================================
section_heading(20, 'COMMERCIAL PROPOSAL')
classification_note('PROJECT-SPECIFIC')

sub_heading('20.1 Pricing Structure')

bold_paragraph('Engagement Type: Fixed-Price Development & Deployment')

doc.add_paragraph()
bold_paragraph('Scope Coverage:')
bullet('Complete development, testing, and deployment of 7 intelligent agents (4 automated, 3 assisted)')
bullet('29 screens + 11 modals with responsive design and WCAG 2.0 AA accessibility')
bullet('66 APIs (58 vendor-built, 7 integrations, 1 orchestration)')
bullet('6-month hypercare support')
bullet('Knowledge transfer and team enablement')
bullet('All team costs and AI tooling (6 Claude Opus subscriptions)')

doc.add_paragraph()
bold_paragraph('Quality Guarantees:')
bullet('80% code coverage (backend and frontend)')
bullet('< 3 second API response time (P95)')
bullet('99.5% system availability')
bullet('90%+ KPI achievement at agent level')

doc.add_paragraph()
doc.add_paragraph(
    'FYI \u2014 Detailed pricing breakdown and payment milestone schedule are provided in the '
    'Pricing Agreement (separate document).'
)

sub_heading('20.2 Milestone Payment Plan')

add_table(
    ['Milestone', 'Description', 'Payment %'],
    [
        ('M0', 'Project Kickoff & Design Approval', '20%'),
        ('M1', 'MVP Development Complete (UAT Ready)', '30%'),
        ('M2', 'MVP Go-Live (1 Live Customer)', '20%'),
        ('M3', 'Phase 2 UAT Complete', '20%'),
        ('M4', 'Full Project Completion & Handover', '10%'),
    ]
)

bold_paragraph('Payment Terms:')
bullet('Net 30 days from invoice date')
bullet('Monthly invoices for ongoing hypercare')
bullet('Late payment penalty: 1.5% per month on overdue amounts')

sub_heading('20.3 Assumptions')
bullet('Pricing based on defined RFP scope (7 agents, 29 screens, 66 APIs, specified NFRs)')
bullet('External services (screening, bank validation, IAM) available and stable')
bullet('ERP integration APIs accessible and documented')
bullet('Required stakeholders available for timely review and approvals')
bullet('Multi-tenant architecture follows agreed design patterns')
bullet('Any deviation from these assumptions may impact cost and timeline')

sub_heading('20.4 Change Billing Model')

bold_paragraph('Baseline Scope: 7 intelligent agents (4 automated, 3 assisted), 29 screens + 11 modals, 66 APIs (58 vendor-built, 7 integrations, 1 orchestration), 25+ core features')

doc.add_paragraph()
bold_paragraph('Change Management Process:')
numbered('Change Request Submission: HighRadius submits change request with business justification')
numbered('Impact Assessment: Metapointer provides detailed analysis within 3 business days (technical feasibility, schedule impact, resource requirements, cost implications)')
numbered('Change Advisory Board (CAB) Review: HighRadius approves, defers, or rejects change')
numbered('Change Order: If approved, formal change order signed with cost and schedule adjustments')
numbered('Billing: Cost adjustment deducted from final payment or invoiced separately per change order')

doc.add_page_break()

# ============================================================
# 21. LEGAL & CONTRACTUAL TERMS
# ============================================================
section_heading(21, 'LEGAL & CONTRACTUAL TERMS')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# 22. ABOUT METAPOINTER
# ============================================================
section_heading(22, 'ABOUT METAPOINTER')
classification_note('STANDARD')
standard_placeholder()

doc.add_page_break()

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Prepared By: ').bold = True
p.add_run('Metapointer\n')
p.add_run('For: ').bold = True
p.add_run('HighRadius Corporation\n')
p.add_run('Date: ').bold = True
p.add_run('February 28, 2026')

# Save
doc.save('Supplier_Portal_Proposal_v3.2.docx')
print("SUCCESS: Created Supplier_Portal_Proposal_v3.2.docx with all 22 sections")
